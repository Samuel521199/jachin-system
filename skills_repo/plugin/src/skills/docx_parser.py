"""Word 文档解析"""
import logging
from pathlib import Path
from typing import Dict, Any, List

logger = logging.getLogger(__name__)


async def parse_docx(path: str) -> Dict[str, Any]:
    """解析 .docx 获取纯文本"""
    try:
        file_path = Path(path).resolve()
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {path}", "text": ""}
        if file_path.suffix.lower() not in (".docx", ".doc"):
            return {"success": False, "error": "仅支持 .docx 文件", "text": ""}

        from docx import Document
        doc = Document(file_path)
        paragraphs: List[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                paragraphs.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        full_text = "\n".join(paragraphs)
        return {"success": True, "text": full_text, "paragraphs": paragraphs}
    except ImportError as e:
        return {"success": False, "error": f"python-docx 未安装: {e}", "text": ""}
    except Exception as e:
        logger.error(f"parse_docx failed: {e}", exc_info=True)
        return {"success": False, "error": str(e), "text": ""}
