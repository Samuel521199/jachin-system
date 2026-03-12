"""
com.jachin.docx-parser - Word 简历解析
解析 .docx 格式的 Word 文档，提取正文文本
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional, List

try:
    from core.skills.base_skill import BaseSkill
except ImportError:
    BaseSkill = object

logger = logging.getLogger(__name__)


class DocxParserSkill(BaseSkill):
    """Word 文档解析技能"""

    def __init__(self, manifest: Dict[str, Any]):
        if BaseSkill is not object:
            super().__init__(manifest)
        else:
            self.manifest = manifest
            self.skill_id = manifest.get("id", "unknown")

    async def execute(self, capability: str, params: Dict[str, Any], context: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        if BaseSkill is not object:
            return await super().execute(capability, params, context)
        if capability == "parse_docx":
            return await self.parse_docx(params)
        return {"success": False, "error": f"Unknown capability: {capability}"}

    async def parse_docx(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """解析 .docx 文件，提取正文文本"""
        try:
            path = params.get("path", "")
            if not path:
                return {"success": False, "error": "path is required"}

            file_path = Path(path).resolve()
            if not file_path.exists():
                return {"success": False, "error": f"File not found: {path}"}
            if file_path.suffix.lower() not in (".docx", ".doc"):
                return {"success": False, "error": "File must be .docx"}

            from docx import Document

            doc = Document(file_path)
            paragraphs: List[str] = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    paragraphs.append(t)

            # 表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        ct = cell.text.strip()
                        if ct:
                            row_text.append(ct)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            full_text = "\n".join(paragraphs)
            return {
                "success": True,
                "text": full_text,
                "paragraphs": paragraphs,
                "chars_extracted": len(full_text),
            }
        except ImportError as e:
            return {"success": False, "error": f"python-docx not installed: {e}"}
        except Exception as e:
            logger.error(f"parse_docx failed: {e}", exc_info=True)
            return {"success": False, "error": str(e)}
