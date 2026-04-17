"""
将 WebSocket / 网关传入的附件实体转为 OpenAI 兼容的 user content（纯文本或 text+image_url 列表）。

- 图片：读本地路径或 base64 → 可选 Pillow 缩放 → data:image/...;base64,...
- 文档：PDF / docx / xlsx / 纯文本（txt、md、csv、log 等）→ 预限流后再 _truncate_doc 拼入 text 段

失败单条附件不拖死整轮：记录日志并跳过该条。
"""
from __future__ import annotations

import base64
import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import unquote

logger = logging.getLogger(__name__)

MAX_DOC_EXTRACT_CHARS = 10_000
# xlsx 在截断为 MAX_DOC_EXTRACT_CHARS 之前，若先拼接整表再截断会长时间占用 CPU/内存（大表可达数十秒）
MAX_XLSX_PRE_TRUNCATE_CHARS = 80_000
MAX_XLSX_ROWS_PER_SHEET = 2_500
MAX_XLSX_SHEETS = 20
# docx：须在 _truncate_doc 之前限流，否则大文档遍历全段落/全表可达数十秒（与 xlsx 同类问题）
MAX_DOCX_PRE_TRUNCATE_CHARS = 80_000
MAX_DOCX_PARAGRAPHS = 4_000
MAX_DOCX_TABLES = 60
MAX_DOCX_ROWS_PER_TABLE = 600
MAX_DOCX_CELLS_PER_ROW = 64
# txt / md / csv / log / 通用 UTF-8 字节解码：先限读再限字，避免超大附件一次读入内存
MAX_PLAIN_TEXT_READ_BYTES = 512_000
MAX_PLAIN_TEXT_PRE_TRUNCATE_CHARS = 80_000
# PDF：多页逐页抽取在无上限时可达数十秒；与纯文本共用字符预算量级
MAX_PDF_PAGES = 120
MAX_PDF_PRE_TRUNCATE_CHARS = 80_000
MAX_IMAGE_READ_BYTES = 5 * 1024 * 1024
MAX_IMAGE_LONG_EDGE = 1536
JPEG_QUALITY = 82

# 含图时追加到 user 文本段：避免 VL 在 ReAct 下误调网页抓取、或否认「能看见图」
_MULTIMODAL_VISION_HINT_ZH = (
    "\n\n【多模态】本条用户消息含上传图片，请直接依据图像作答（含图中文字/OCR）；"
    "不要声称无法读图或缺少图像识别能力。"
    "请勿仅因会话历史中出现 http(s) 链接或旧轮 Observation 就调用网页抓取；"
    "历史里的新浪/新闻正文与当前截图无关。除非用户在本轮明确给出要你抓取的 URL。"
)

_EXT_MIME = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/plain",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
}


def _display_name(att: dict[str, Any]) -> str:
    return str(att.get("name") or att.get("filename") or "attachment")[:256]


def _coerce_remote_image_url(att: dict[str, Any]) -> str:
    """
    网关 / 客户端可能传 OpenAI 风格：
    - image_url: \"https://...\" 或 \"data:image/jpeg;base64,...\"
    - image_url: {\"url\": \"...\"}（若为 dict 则取内层 url，禁止 str(dict) 误用）
    """
    v = att.get("image_url")
    if isinstance(v, dict):
        u = str(v.get("url") or "").strip()
        if u:
            return u
    elif isinstance(v, str) and v.strip():
        return v.strip()
    u2 = att.get("url")
    if isinstance(u2, dict):
        u = str(u2.get("url") or "").strip()
        if u:
            return u
    elif isinstance(u2, str) and u2.strip():
        return u2.strip()
    return ""


# 纯 Base64、无 local_path 时须从文件名推断类型，否则 .docx 会误判为纯文本而把 ZIP 当 UTF-8 解码成乱码（PK…）
_DOC_EXT_ORDER = (
    ".docx",
    ".xlsx",
    ".xls",
    ".doc",
    ".pdf",
    ".txt",
    ".md",
    ".csv",
    ".log",
)


def _infer_attachment_suffix(path: Path | None, display_name: str, mime: str) -> str:
    if path is not None:
        return path.suffix.lower()
    low = (display_name or "").lower().strip()
    for ext in _DOC_EXT_ORDER:
        if low.endswith(ext):
            return ext
    m = (mime or "").lower()
    if "pdf" in m or m == "application/pdf":
        return ".pdf"
    if "wordprocessingml" in m and "spreadsheetml" not in m:
        return ".docx"
    if "msword" in m:
        return ".docx"
    if "spreadsheetml.sheet" in m:
        return ".xlsx"
    if m == "application/vnd.ms-excel":
        return ".xls"
    if m.startswith("text/"):
        return ".txt"
    return ""


def _guess_mime(att: dict[str, Any], path: Path | None) -> str:
    m = str(att.get("mime") or att.get("content_type") or "").strip().lower()
    # 浏览器对 .docx 常报 application/octet-stream，须结合后缀推断，否则后续易误判
    _generic_octet = m in ("application/octet-stream", "binary/octet-stream", "application/x-msdownload")
    if m and not _generic_octet:
        return m[:128]
    if path is not None:
        pm = _EXT_MIME.get(path.suffix.lower(), "")
        if pm:
            return pm
    name = str(att.get("name") or att.get("filename") or "")
    inferred = _EXT_MIME.get(_infer_attachment_suffix(None, name, m), "")
    if inferred:
        return inferred
    return m[:128] if m else ""


def _resolve_local_path(att: dict[str, Any]) -> Path | None:
    for k in ("local_path", "file_path", "path", "absolute_path"):
        v = att.get(k)
        if isinstance(v, str) and v.strip():
            p = Path(v.strip()).expanduser()
            try:
                if p.is_file():
                    return p.resolve()
            except OSError:
                continue
    uri = str(att.get("uri") or "").strip()
    if uri.lower().startswith("file://"):
        raw = uri[7:]
        if raw.startswith("/") and len(raw) >= 4 and raw[2] == ":":
            raw = raw[1:]
        elif raw.startswith("//"):
            raw = raw[2:]
        try:
            p = Path(unquote(raw)).expanduser()
            if p.is_file():
                return p.resolve()
        except OSError:
            pass
    return None


def _decode_data_image_url(data_url: str) -> tuple[bytes | None, str]:
    """解析 data:image/...;base64,... 为原始字节与 MIME（失败返回 None, \"\"）。"""
    s = (data_url or "").strip()
    if not s.lower().startswith("data:image/"):
        return None, ""
    try:
        comma = s.index(",")
    except ValueError:
        return None, ""
    header = s[:comma]
    rest = s[comma + 1 :]
    if ";base64" not in header.lower():
        return None, ""
    raw = base64.b64decode(re.sub(r"\s+", "", rest), validate=False)
    mime = "image/jpeg"
    m = re.match(r"data:([^;]+)", header, re.I)
    if m:
        mime = (m.group(1) or "image/jpeg").strip() or "image/jpeg"
    return raw, mime


def _load_raw_bytes(att: dict[str, Any]) -> tuple[bytes | None, Path | None, str]:
    """返回 (bytes 或 None, 可选路径, mime 提示)。"""
    mime = _guess_mime(att, None)
    # 顶层字符串字段上的 data URL（部分网关透传）
    for uk in ("image_url", "url"):
        u = att.get(uk)
        if isinstance(u, str) and u.strip().lower().startswith("data:image/"):
            raw, dm = _decode_data_image_url(u.strip())
            if raw is not None:
                return raw, None, dm or mime
    for fk in ("base64_data", "data_base64", "base64"):
        b64 = att.get(fk)
        if isinstance(b64, str) and b64.strip():
            try:
                raw = base64.b64decode(re.sub(r"\s+", "", b64.strip()), validate=False)
                return raw, None, mime
            except Exception as e:
                logger.debug("[multimodal] base64 解码失败 name=%s err=%s", _display_name(att), e)
    p = _resolve_local_path(att)
    if p is not None:
        mime = _guess_mime(att, p)
        try:
            sz = p.stat().st_size
            if sz > MAX_IMAGE_READ_BYTES:
                logger.warning(
                    "[multimodal] 附件过大已跳过 name=%s bytes=%s max=%s",
                    _display_name(att),
                    sz,
                    MAX_IMAGE_READ_BYTES,
                )
                return None, p, mime
            return p.read_bytes(), p, mime
        except OSError as e:
            logger.debug("[multimodal] 读文件失败 path=%s err=%s", p, e)
    return None, None, mime


def _maybe_resize_image(data: bytes, mime_guess: str) -> tuple[bytes, str]:
    try:
        from PIL import Image
    except ImportError:
        return data, mime_guess or "image/jpeg"
    try:
        im = Image.open(BytesIO(data))
        im = im.convert("RGB")
        w, h = im.size
        mx = max(w, h)
        if mx > MAX_IMAGE_LONG_EDGE:
            scale = MAX_IMAGE_LONG_EDGE / float(mx)
            im = im.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)
        out = BytesIO()
        im.save(out, format="JPEG", quality=JPEG_QUALITY, optimize=True)
        return out.getvalue(), "image/jpeg"
    except Exception as e:
        logger.debug("[multimodal] 图片压缩跳过 err=%s", e)
        return data, mime_guess or "image/jpeg"


def _truncate_doc(s: str) -> str:
    s = (s or "").strip()
    if len(s) <= MAX_DOC_EXTRACT_CHARS:
        return s
    return s[: MAX_DOC_EXTRACT_CHARS - 20] + "\n…(已截断)"


_PLAIN_BUDGET_TAIL = "\n…(文本体量过大，仅提取前段；需要全量请用工具读文件)…"


def _budget_unicode_text(
    s: str,
    max_chars: int,
    *,
    tail: str = _PLAIN_BUDGET_TAIL,
) -> str:
    """在写入 prompt 前限制字符串体量，避免超长文本占用 CPU/内存（与 xlsx 预截断同策略）。"""
    s = s or ""
    if len(s) <= max_chars:
        return s
    room = max_chars - len(tail) - 1
    if room < 32:
        return s[:max_chars]
    return s[:room].rstrip() + tail


def _extract_pdf_text(path: Path | None, data: bytes | None) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        fitz = None  # type: ignore[misc, assignment]
    if fitz is not None:
        try:
            doc = (
                fitz.open(str(path))
                if path is not None
                else fitz.open(stream=data, filetype="pdf")
            )
            try:
                parts: list[str] = []
                acc_len = 0
                n_pages = len(doc)
                last_i = -1
                char_capped = False
                for i in range(n_pages):
                    if i >= MAX_PDF_PAGES:
                        parts.append(f"(PDF 仅解析前 {MAX_PDF_PAGES} 页，共 {n_pages} 页…)")
                        break
                    t = doc[i].get_text() or ""
                    parts.append(t)
                    acc_len += len(t) + 1
                    last_i = i
                    if acc_len >= MAX_PDF_PRE_TRUNCATE_CHARS:
                        parts.append("…(PDF 正文过长，仅提取前段…)")
                        char_capped = True
                        break
                if char_capped and last_i >= 0 and last_i < n_pages - 1:
                    parts.append(f"(另有 {n_pages - last_i - 1} 页 PDF 未读…)")
                return _budget_unicode_text("\n".join(parts), MAX_PDF_PRE_TRUNCATE_CHARS)
            finally:
                doc.close()
        except Exception as e:
            logger.debug("[multimodal] PyMuPDF 解析失败 err=%s", e)
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        bio = BytesIO(data) if data is not None else None
        reader = PdfReader(bio if bio is not None else str(path))
        parts: list[str] = []
        acc_len = 0
        n_pdf = len(reader.pages)
        char_capped = False
        last_pi = -1
        for pi, page in enumerate(reader.pages):
            if pi >= MAX_PDF_PAGES:
                parts.append(f"(PDF 仅解析前 {MAX_PDF_PAGES} 页，共 {n_pdf} 页…)")
                break
            t = page.extract_text() or ""
            parts.append(t)
            acc_len += len(t) + 1
            last_pi = pi
            if acc_len >= MAX_PDF_PRE_TRUNCATE_CHARS:
                parts.append("…(PDF 正文过长，仅提取前段…)")
                char_capped = True
                break
        if char_capped and last_pi >= 0 and last_pi < n_pdf - 1:
            parts.append(f"(另有 {n_pdf - last_pi - 1} 页 PDF 未读…)")
        return _budget_unicode_text("\n".join(parts), MAX_PDF_PRE_TRUNCATE_CHARS)
    except Exception as e:
        logger.debug("[multimodal] pypdf 解析失败 err=%s", e)
        return ""


def _extract_docx_text(path: Path | None, data: bytes | None) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        if path is not None:
            d = Document(str(path))
        elif data is not None:
            d = Document(BytesIO(data))
        else:
            return ""
        parts: list[str] = []
        acc_len = 0
        npara = 0
        for p in d.paragraphs:
            npara += 1
            if npara > MAX_DOCX_PARAGRAPHS:
                parts.append("…(段落数量过多已省略…)")
                break
            t = (p.text or "").strip()
            if t:
                parts.append(t)
                acc_len += len(t) + 1
            if acc_len >= MAX_DOCX_PRE_TRUNCATE_CHARS:
                parts.append("…(Word 正文过长，仅提取前段供模型参考；需要全量请用工具读文件)…")
                break
        nt = 0
        if acc_len < MAX_DOCX_PRE_TRUNCATE_CHARS:
            for table in d.tables:
                nt += 1
                if nt > MAX_DOCX_TABLES:
                    parts.append("…(表格数量过多已省略…)")
                    break
                for ri, row in enumerate(table.rows):
                    if ri >= MAX_DOCX_ROWS_PER_TABLE:
                        parts.append("(本表行数过多已省略…)")
                        break
                    cells: list[str] = []
                    for ci, cell in enumerate(row.cells):
                        if ci >= MAX_DOCX_CELLS_PER_ROW:
                            cells.append("…")
                            break
                        ct = (cell.text or "").strip()
                        cells.append(ct)
                    line = "\t".join(cells).rstrip()
                    if line:
                        parts.append(line)
                        acc_len += len(line) + 1
                    if acc_len >= MAX_DOCX_PRE_TRUNCATE_CHARS:
                        parts.append("…(表格内容过长，已截断…)")
                        break
                if acc_len >= MAX_DOCX_PRE_TRUNCATE_CHARS:
                    break
        return "\n".join(parts)
    except Exception as e:
        logger.debug("[multimodal] docx 解析失败 err=%s", e)
        return ""


def _extract_xlsx_text(path: Path | None, data: bytes | None) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ""
    try:
        if path is not None:
            wb = load_workbook(filename=str(path), read_only=True, data_only=True)
        elif data is not None:
            wb = load_workbook(filename=BytesIO(data), read_only=True, data_only=True)
        else:
            return ""
        parts: list[str] = []
        acc_len = 0
        for si, ws in enumerate(wb.worksheets):
            if si >= MAX_XLSX_SHEETS:
                parts.append("(另有工作表已省略…)")
                break
            parts.append(f"=== Sheet: {ws.title} ===")
            acc_len += len(parts[-1]) + 1
            row_i = 0
            for row in ws.iter_rows(values_only=True):
                row_i += 1
                if row_i > MAX_XLSX_ROWS_PER_SHEET:
                    parts.append("(本表行数过多已省略…)")
                    acc_len += 24
                    break
                cells: list[str] = []
                for v in row:
                    if v is None:
                        cells.append("")
                    else:
                        cells.append(str(v).strip())
                line = "\t".join(cells).rstrip()
                if line:
                    parts.append(line)
                    acc_len += len(line) + 1
                if acc_len >= MAX_XLSX_PRE_TRUNCATE_CHARS:
                    parts.append("…(表格体量过大，仅提取前段供模型参考；需要全量请用工具读文件或拆分上传)")
                    break
            if acc_len >= MAX_XLSX_PRE_TRUNCATE_CHARS:
                break
        try:
            wb.close()
        except Exception:
            pass
        return "\n".join(parts)
    except Exception as e:
        logger.debug("[multimodal] xlsx 解析失败 err=%s", e)
        return ""


def _extract_plain_text(path: Path | None, data: bytes | None) -> str:
    """
    .txt / .md / .csv / .log 及回退路径下的原始字节 UTF-8 解码。
    先限制读取字节数再限制字符数，避免单附件内嵌超大日志/CSV 拖慢 to_thread。
    """
    raw: bytes | None = None
    try:
        if data is not None:
            raw = data[:MAX_PLAIN_TEXT_READ_BYTES]
        elif path is not None:
            with path.open("rb") as f:
                raw = f.read(MAX_PLAIN_TEXT_READ_BYTES)
    except OSError:
        return ""
    if raw is None:
        return ""
    text = raw.decode("utf-8", errors="replace")
    return _budget_unicode_text(text, MAX_PLAIN_TEXT_PRE_TRUNCATE_CHARS)


def _is_image_mime(mime: str, path: Path | None, att: dict[str, Any]) -> bool:
    if (mime or "").lower().startswith("image/"):
        return True
    if bool(att.get("has_image") or att.get("is_image")):
        return True
    if path is not None and path.suffix.lower() in (
        ".jpg",
        ".jpeg",
        ".png",
        ".gif",
        ".webp",
        ".bmp",
    ):
        return True
    return False


def _is_document(path: Path | None, mime: str, display_name: str = "") -> bool:
    m = (mime or "").lower()
    if m.startswith("text/"):
        return True
    if "pdf" in m or m == "application/pdf":
        return True
    if "wordprocessingml" in m or "msword" in m:
        return True
    if "spreadsheetml" in m or m == "application/vnd.ms-excel":
        return True
    suf = _infer_attachment_suffix(path, display_name, mime)
    return suf in (".txt", ".md", ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".log")


def build_openai_user_content(user_text: str, attachments: list[dict[str, Any]]) -> str | list[dict[str, Any]]:
    """
    无图片时返回纯字符串（兼容旧 ReAct / 日志）；含图片时返回 OpenAI 多模态 content 列表。
    """
    if not attachments:
        return user_text or ""

    doc_chunks: list[str] = []
    image_parts: list[dict[str, Any]] = []

    for att in attachments:
        if not isinstance(att, dict):
            continue
        name = _display_name(att)
        # 远端 URL 或 OpenAI 风格嵌套 {\"url\": \"...\"}；禁止 str(dict) 误解析
        _remote = _coerce_remote_image_url(att)
        if _remote.startswith("http://") or _remote.startswith("https://"):
            if bool(att.get("has_image") or att.get("is_image")) or _guess_mime(att, None).startswith(
                "image/"
            ):
                image_parts.append({"type": "image_url", "image_url": {"url": _remote}})
                continue
        # data:image/...;base64,...（桌面 / 网关透传常见；此前仅认 http(s)，导致图片被静默丢弃）
        if _remote.lower().startswith("data:image/"):
            _raw_chk, _ = _decode_data_image_url(_remote)
            if _raw_chk is None:
                logger.debug("[multimodal] data URL 无法解析，跳过 name=%s", name)
                continue
            if len(_raw_chk) > MAX_IMAGE_READ_BYTES:
                logger.warning("[multimodal] data URL 图片过大已跳过 name=%s", name)
                continue
            image_parts.append({"type": "image_url", "image_url": {"url": _remote}})
            continue
        raw, path, mime = _load_raw_bytes(att)

        if raw is None and path is None:
            logger.debug("[multimodal] 无可用实体，跳过元数据-only 附件 name=%s", name)
            continue

        if raw is not None and len(raw) > MAX_IMAGE_READ_BYTES:
            logger.warning("[multimodal] 附件过大已跳过 name=%s", name)
            continue

        if _is_image_mime(mime, path, att):
            data = raw if raw is not None else (path.read_bytes() if path else None)
            if not data:
                continue
            data, out_mime = _maybe_resize_image(data, mime)
            b64 = base64.b64encode(data).decode("ascii")
            url = f"data:{out_mime};base64,{b64}"
            image_parts.append({"type": "image_url", "image_url": {"url": url}})
            continue

        if _is_document(path, mime, name):
            suf = _infer_attachment_suffix(path, name, mime)
            extracted = ""
            if suf == ".pdf" or "pdf" in (mime or "").lower():
                extracted = _truncate_doc(_extract_pdf_text(path, raw))
            elif suf == ".docx":
                extracted = _truncate_doc(_extract_docx_text(path, raw))
            elif suf == ".xlsx":
                extracted = _truncate_doc(_extract_xlsx_text(path, raw))
            elif suf == ".xls":
                logger.warning(
                    "[multimodal] .xls 为旧版 Excel 二进制格式，openpyxl 不支持；请另存为 .xlsx name=%s",
                    name,
                )
                extracted = _truncate_doc(
                    "（本文件为 .xls 旧格式，无法自动提取表格文本；请另存为 .xlsx 后重新上传。）"
                )
            elif suf == ".doc":
                logger.warning(
                    "[multimodal] .doc 旧版 Word 需本机 Word/LibreOffice 转换，已跳过正文提取 name=%s",
                    name,
                )
                extracted = ""
            else:
                extracted = _truncate_doc(_extract_plain_text(path, raw))
            if extracted.strip():
                doc_chunks.append(f"[附件: {name} 内容]\n\n{extracted}")
            else:
                doc_chunks.append(f"[附件: {name} 内容]\n\n(未能提取文本或文件为空)")
        else:
            logger.debug("[multimodal] 不支持的类型已跳过 mime=%s name=%s", mime, name)

    doc_preamble = "\n\n".join(doc_chunks).strip()
    body = (user_text or "").strip()
    if doc_preamble and body:
        combined_text = doc_preamble + "\n\n" + body
    elif doc_preamble:
        combined_text = doc_preamble
    else:
        combined_text = body

    if not image_parts:
        return combined_text

    if not combined_text:
        combined_text = "（用户未附带文字说明）"

    combined_text = combined_text + _MULTIMODAL_VISION_HINT_ZH

    # 与 DashScope 官方 MultiModalConversation 示例一致：先图后文（避免兼容层忽略 image_url）
    return [*image_parts, {"type": "text", "text": combined_text}]
