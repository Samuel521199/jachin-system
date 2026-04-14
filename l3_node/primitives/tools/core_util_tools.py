"""
Jachin L3 — 原生轻量实用工具 (util:* / sys:*)

供大模型补齐：绝对时间、安全算术、编解码、轻量网络与主机状态等。
所有 run_* 返回可 JSON 序列化的 dict；异常时 {"ok": False, "error": ...}。
"""
from __future__ import annotations

import ast
import base64
import difflib
import hashlib
import json
import math
import os
import platform
import re
import socket
import subprocess
import time
import uuid
from datetime import datetime, timedelta, timezone
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.parse import quote
from urllib.request import Request, urlopen

# ---------------------------------------------------------------------------
# 小工具：统一成功 / 失败包装
# ---------------------------------------------------------------------------


def _ok(result: Any) -> dict[str, Any]:
    return {"ok": True, "result": result}


def _err(msg: str) -> dict[str, Any]:
    return {"ok": False, "error": str(msg)}


# ---------------------------------------------------------------------------
# 类别一：时间与数学
# ---------------------------------------------------------------------------


def run_datetime_calc(**kwargs: Any) -> dict[str, Any]:
    """util:datetime_calc — base_time 可选 ISO8601；add_days；target_timezone（IANA）。"""
    try:
        from zoneinfo import ZoneInfo

        add_days = int(kwargs.get("add_days", 0))
        tz_name = str(kwargs.get("target_timezone") or "Asia/Shanghai").strip() or "UTC"
        try:
            tz = ZoneInfo(tz_name)
        except Exception:
            tz = timezone.utc
            tz_name = "UTC"

        base_raw = kwargs.get("base_time")
        if base_raw is None or str(base_raw).strip() == "":
            dt = datetime.now(tz=tz)
        else:
            s = str(base_raw).strip()
            if s.endswith("Z"):
                s = s[:-1] + "+00:00"
            dt = datetime.fromisoformat(s)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=tz)
            else:
                dt = dt.astimezone(tz)

        out = dt + timedelta(days=add_days)
        return _ok(
            {
                "iso_local": out.isoformat(),
                "timezone": tz_name,
                "add_days": add_days,
                "unix_timestamp": out.timestamp(),
            }
        )
    except Exception as e:
        return _err(e)


def _cron_field_zh(field: str, idx: int) -> str:
    """极简五段式 cron 字段中文描述（启发式）。"""
    f = field.strip()
    names = ("分", "时", "日", "月", "周")
    label = names[idx] if idx < len(names) else f"段{idx}"
    if f == "*":
        return f"{label}：每分钟/每次" if idx == 0 else f"{label}：任意"
    if f.isdigit():
        return f"{label}：{f}"
    if f.startswith("*/"):
        return f"{label}：每 {f[2:]} 单位"
    return f"{label}：{f}"


def _cron_next_three_via_croniter(expr: str) -> list[str]:
    from croniter import croniter

    base = datetime.now(timezone.utc)
    it = croniter(expr, base)
    next_times: list[str] = []
    for _ in range(3):
        nxt = it.get_next(datetime)
        if nxt.tzinfo is None:
            nxt = nxt.replace(tzinfo=timezone.utc)
        next_times.append(nxt.astimezone(timezone.utc).isoformat())
    return next_times


def _cron_next_three_via_apscheduler(expr: str) -> list[str]:
    """5 段 Unix cron；依赖项目已有的 apscheduler（无需 croniter）。"""
    from apscheduler.triggers.cron import CronTrigger

    ct = CronTrigger.from_crontab(expr)
    prev = None
    now = datetime.now(timezone.utc)
    next_times: list[str] = []
    for _ in range(3):
        nxt = ct.get_next_fire_time(prev, now)
        if nxt is None:
            break
        next_times.append(nxt.astimezone(timezone.utc).isoformat())
        prev = nxt
        now = nxt + timedelta(microseconds=1)
    return next_times


def run_cron_explain(**kwargs: Any) -> dict[str, Any]:
    """util:cron_explain — 中文简述 + 接下来 3 次触发。5 段可用 APScheduler；6 段（含秒）优先 croniter。"""
    try:
        expr = str(kwargs.get("cron_expr") or kwargs.get("expression") or "").strip()
        if not expr:
            return _err("cron_expr 不能为空")

        parts = expr.split()
        if len(parts) not in (5, 6):
            return _err("Cron 应为 5 段（分 时 日 月 周）或 6 段（秒 分 时 日 月 周）")

        zh_parts = [_cron_field_zh(p, i) for i, p in enumerate(parts[-5:])]
        summary_zh = "；".join(zh_parts)

        next_times: list[str] = []
        engine = ""

        if len(parts) == 6:
            try:
                next_times = _cron_next_three_via_croniter(expr)
                engine = "croniter"
            except ImportError:
                return _err(
                    "6 段 Cron（含秒）需要安装 croniter。若清华等镜像暂无该包，请改用官方源："
                    "pip install croniter -i https://pypi.org/simple；或改为 5 段 Unix cron（不含秒）。"
                )
        else:
            # 5 段：优先 APScheduler（项目已依赖），不强制 croniter（避免镜像缺包）
            try:
                next_times = _cron_next_three_via_apscheduler(expr)
                engine = "apscheduler"
            except Exception:
                try:
                    next_times = _cron_next_three_via_croniter(expr)
                    engine = "croniter"
                except ImportError:
                    return _err(
                        "无法解析该 Cron。请安装 croniter（官方源）：pip install croniter -i https://pypi.org/simple"
                    )

        return _ok(
            {
                "cron_expr": expr,
                "engine": engine,
                "summary_zh": f"表达式「{expr}」大致含义：{summary_zh}（启发式说明，复杂规则请以实际调度为准）",
                "next_three_iso_utc": next_times,
            }
        )
    except Exception as e:
        return _err(e)


class _SafeMathVisitor(ast.NodeVisitor):
    """仅允许 + - * / 与一元正负、数值常量（Decimal）。"""

    def visit(self, node: ast.AST) -> Decimal:  # type: ignore[override]
        if isinstance(node, ast.Expression):
            return self.visit(node.body)
        if isinstance(node, ast.Constant):
            if isinstance(node.value, (int, float)):
                return Decimal(str(node.value))
            if isinstance(node.value, str):
                try:
                    return Decimal(node.value.strip())
                except InvalidOperation as e:
                    raise ValueError(f"无法将字符串转为数字: {node.value!r}") from e
            raise ValueError("不允许的常量类型")
        if isinstance(node, ast.Num):  # py<3.8 兼容
            return Decimal(str(node.n))
        if isinstance(node, ast.UnaryOp) and isinstance(node.op, (ast.USub, ast.UAdd)):
            v = self.visit(node.operand)
            return -v if isinstance(node.op, ast.USub) else v
        if isinstance(node, ast.BinOp) and isinstance(node.op, (ast.Add, ast.Sub, ast.Mult, ast.Div)):
            left = self.visit(node.left)
            right = self.visit(node.right)
            if isinstance(node.op, ast.Add):
                return left + right
            if isinstance(node.op, ast.Sub):
                return left - right
            if isinstance(node.op, ast.Mult):
                return left * right
            if isinstance(node.op, ast.Div):
                if right == 0:
                    raise ZeroDivisionError("除数为零")
                return left / right
        raise ValueError(f"不允许的语法节点: {type(node).__name__}")


def run_precise_math(**kwargs: Any) -> dict[str, Any]:
    """util:precise_math — 安全四则运算，Decimal，禁止 eval。"""
    try:
        expression = str(kwargs.get("expression") or "").strip()
        if not expression:
            return _err("expression 不能为空")
        tree = ast.parse(expression, mode="eval")
        val: Decimal = _SafeMathVisitor().visit(tree)
        # 统一字符串输出，避免 JSON 对 Decimal 不友好
        return _ok({"decimal_str": format(val, "f"), "repr": str(val)})
    except Exception as e:
        return _err(e)


# ---------------------------------------------------------------------------
# 类别二：文本与密码学
# ---------------------------------------------------------------------------


def run_uuid_gen(**kwargs: Any) -> dict[str, Any]:
    """util:uuid_gen — 生成 uuid4。"""
    try:
        _ = kwargs  # 接受任意 kwargs，忽略
        return _ok({"uuid": str(uuid.uuid4())})
    except Exception as e:
        return _err(e)


def run_hash_crypto(**kwargs: Any) -> dict[str, Any]:
    """util:hash_crypto — algo: md5 | sha256 | base64_encode | base64_decode"""
    try:
        text = kwargs.get("text")
        if text is None:
            return _err("text 不能为空")
        algo = str(kwargs.get("algo") or "sha256").strip().lower()
        raw: bytes
        if isinstance(text, bytes):
            raw = text
        else:
            s = str(text)
            if algo == "base64_decode":
                raw = base64.b64decode(s.encode("ascii"), validate=False)
                return _ok({"algo": algo, "output_text": raw.decode("utf-8", errors="replace")})
            raw = s.encode("utf-8")

        if algo == "md5":
            return _ok({"algo": algo, "hex": hashlib.md5(raw).hexdigest()})
        if algo == "sha256":
            return _ok({"algo": algo, "hex": hashlib.sha256(raw).hexdigest()})
        if algo == "base64_encode":
            return _ok({"algo": algo, "b64": base64.b64encode(raw).decode("ascii")})
        return _err(f"不支持的 algo: {algo}")
    except Exception as e:
        return _err(e)


def _normalize_json_path(path: str) -> str:
    """将 a[0].b 转为 a.0.b 供按段解析。"""
    s = path.strip()
    s = re.sub(r"\[(\d+)\]", r".\1", s)
    s = re.sub(r"\.\.+", ".", s)
    return s.strip(".")


def _json_get_path(obj: Any, path: str) -> Any:
    cur = obj
    path = _normalize_json_path(path)
    if not path:
        return cur
    for part in path.split("."):
        part = part.strip()
        if not part:
            continue
        if part.isdigit():
            cur = cur[int(part)]
        else:
            cur = cur[part]
    return cur


def run_json_jq(**kwargs: Any) -> dict[str, Any]:
    """util:json_jq — 轻量路径提取；path 支持点号与 [0] 下标。"""
    try:
        json_string = str(kwargs.get("json_string") or "")
        path = str(kwargs.get("path") or "").strip()
        if not json_string:
            return _err("json_string 不能为空")
        if not path:
            return _err("path 不能为空")
        data = json.loads(json_string)
        hit = _json_get_path(data, path)
        return _ok({"path": path, "value": hit})
    except Exception as e:
        return _err(e)


def run_regex_test(**kwargs: Any) -> dict[str, Any]:
    """util:regex_test — pattern + test_cases（字符串列表）。"""
    try:
        pattern = str(kwargs.get("pattern") or "")
        if not pattern:
            return _err("pattern 不能为空")
        cases = kwargs.get("test_cases")
        if cases is None:
            return _err("test_cases 不能为空（JSON 数组）")
        if isinstance(cases, str):
            cases = json.loads(cases)
        if not isinstance(cases, list):
            return _err("test_cases 必须为列表")
        rx = re.compile(pattern)
        out: list[dict[str, Any]] = []
        for c in cases:
            s = str(c)
            m = rx.search(s)
            if not m:
                out.append({"text": s, "matched": False, "groups": []})
            else:
                g = list(m.groups()) if m.lastindex else []
                out.append(
                    {
                        "text": s,
                        "matched": True,
                        "full_match": m.group(0),
                        "groups": g,
                    }
                )
        return _ok({"pattern": pattern, "cases": out})
    except Exception as e:
        return _err(e)


# ---------------------------------------------------------------------------
# 类别三：轻量感知
# ---------------------------------------------------------------------------

_STEALTH_TEXT_MAX = 500_000
_STEALTH_HTML_MAX = 300_000

# 轻装阶段若正文出现下列片段，视为被常见 WAF/CF 拦截页，转入重装旁路
_STEALTH_WAF_MARKERS = (
    "just a moment",
    "enable javascript and cookies to continue",
)

_SIDECAR_UNREACHABLE_MSG = (
    "轻装抓取被拦截，且重装旁路服务未启动（连接被拒绝）。"
    "请在独立终端运行 `uvicorn server:app` 以开启重装刺客服务。"
)


def _stealth_clip(s: str, max_len: int) -> str:
    if len(s) <= max_len:
        return s
    return s[:max_len] + "\n...[truncated]"


def _html_to_naive_plain_text(html: str) -> str:
    """无 lxml 时的极简可见文本（降级路径）。"""
    s = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", html)
    s = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", s)
    s = re.sub(r"<[^>]+>", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _stealth_body_looks_waf_blocked(blob: str) -> bool:
    low = (blob or "").lower()
    return any(m in low for m in _STEALTH_WAF_MARKERS)


def _stealth_pack_from_scrapling_response(resp: Any) -> dict[str, Any]:
    status = int(getattr(resp, "status", 200) or 200)
    try:
        raw_text = resp.get_all_text(strip=True, separator="\n")
        text = str(raw_text) if raw_text is not None else ""
    except Exception:
        text = str(getattr(resp, "text", "") or "")
    try:
        hc = resp.html_content
        html_full = str(hc) if hc is not None else ""
    except Exception:
        html_full = ""
    return {
        "text": _stealth_clip(text, _STEALTH_TEXT_MAX),
        "html_excerpt": _stealth_clip(html_full, _STEALTH_HTML_MAX),
        "http_status": status,
    }


def _stealth_try_inprocess_fast(url: str) -> tuple[dict[str, Any] | None, Exception | None]:
    """
    轻装：优先 curl_cffi（显式 impersonate=chrome），失败再试 scrapling Fetcher。
    返回 (payload, None) 表示拿到 HTTP 响应体； (None, exc) 表示两路均未成功。
    """
    try:
        to = float(os.environ.get("JACHIN_STEALTH_INPROCESS_TIMEOUT", "15"))
    except ValueError:
        to = 15.0

    last_exc: Exception | None = None

    try:
        from curl_cffi import requests as curl_req

        r = curl_req.get(url, impersonate="chrome", timeout=to)
        status = int(getattr(r, "status_code", 200) or 200)
        html = str(r.text or "")
        plain = _html_to_naive_plain_text(html)
        return (
            {
                "text": _stealth_clip(plain, _STEALTH_TEXT_MAX),
                "html_excerpt": _stealth_clip(html, _STEALTH_HTML_MAX),
                "http_status": status,
            },
            None,
        )
    except ImportError:
        pass
    except Exception as e:
        last_exc = e

    try:
        from scrapling.fetchers import Fetcher

        resp = Fetcher.get(url, timeout=to, follow_redirects="safe")
        return (_stealth_pack_from_scrapling_response(resp), None)
    except ImportError as e:
        last_exc = last_exc or e
    except Exception as e:
        last_exc = e

    if last_exc is not None:
        return (None, last_exc)
    return (None, RuntimeError("未安装 curl_cffi 或 scrapling，无法进行轻装抓取"))


def _stealth_sidecar_healthcheck(base: str, requests_mod: Any) -> bool:
    """极短超时探测旁路是否在线（任意 HTTP 响应即视为可达）。"""
    try:
        hc_to = float(os.environ.get("JACHIN_SCRAPLING_HEALTH_TIMEOUT", "1.5"))
    except ValueError:
        hc_to = 1.5
    try:
        requests_mod.get(f"{base}/", timeout=hc_to)
        return True
    except Exception:
        return False


def run_http_ping(**kwargs: Any) -> dict[str, Any]:
    """util:http_ping — GET/HEAD 探测，超时 3s。"""
    try:
        url = str(kwargs.get("url") or "").strip()
        if not url:
            return _err("url 不能为空")
        method = str(kwargs.get("method") or "HEAD").upper()
        if method not in ("HEAD", "GET"):
            method = "HEAD"
        t0 = time.perf_counter()
        req = Request(url, method=method, headers={"User-Agent": "Jachin-util-http_ping/1.0"})
        try:
            with urlopen(req, timeout=3) as resp:
                code = getattr(resp, "status", None) or resp.getcode()
        except HTTPError as he:
            code = he.code
        elapsed_ms = int((time.perf_counter() - t0) * 1000)
        return _ok({"url": url, "method": method, "status_code": code, "elapsed_ms": elapsed_ms})
    except URLError as e:
        return _err(f"网络错误: {e}")
    except Exception as e:
        return _err(e)


def run_stealth_extract(**kwargs: Any) -> dict[str, Any]:
    print("\n" + "🔥" * 20, flush=True)
    print("🚀 [ROUTING ALERT] 最新版智能路由 run_stealth_extract 已成功激活！", flush=True)
    print("🔥" * 20 + "\n", flush=True)
    # [TRACER] 若日志无以上三行：非本文件版本 / PYTHONPATH 指向旧副本 / 进程未重启。
    try:
        try:
            import requests
        except ImportError:
            return _err("未安装 requests，请执行: pip install requests")

        url = str(kwargs.get("url") or "").strip()
        if not url:
            return _err("url 不能为空")

        try:
            http_timeout = float(os.environ.get("JACHIN_SCRAPLING_HTTP_TIMEOUT", "15"))
        except ValueError:
            http_timeout = 15.0

        skip_sidecar = str(os.environ.get("JACHIN_STEALTH_EXTRACT_SKIP_SIDECAR", "")).strip() == "1"
        base = str(os.environ.get("JACHIN_SCRAPLING_SERVICE_BASE", "http://127.0.0.1:8000")).rstrip("/")
        endpoint = f"{base}/api/scrape"

        fast_res, fast_err = _stealth_try_inprocess_fast(url)

        def _ok_fast(payload: dict[str, Any]) -> dict[str, Any]:
            print(
                f"👉 路由判定：in_process_fast, HTTP Status: {payload.get('http_status')}",
                flush=True,
            )
            inner: dict[str, Any] = {
                "url": url,
                "content": {
                    "text": payload.get("text", ""),
                    "html_excerpt": payload.get("html_excerpt", ""),
                    "http_status": payload.get("http_status"),
                    "via": "in_process_fast",
                },
            }
            assert inner["content"].get("via") == "in_process_fast"
            return _ok(inner)

        if fast_res is not None:
            st = int(fast_res["http_status"])
            blob = (fast_res.get("html_excerpt") or "") + "\n" + (fast_res.get("text") or "")
            blocked = _stealth_body_looks_waf_blocked(blob)
            if st == 200 and not blocked:
                return _ok_fast(fast_res)
            if st not in (403, 503) and not blocked:
                return _ok_fast(fast_res)

        # 未在轻装阶段 return：fast_res is None，或 403/503 / WAF 拦截页，需重装旁路
        if skip_sidecar:
            bits: list[str] = [
                "需重装旁路但已设置 JACHIN_STEALTH_EXTRACT_SKIP_SIDECAR=1（仅轻装阶段允许）。",
            ]
            if fast_err is not None:
                bits.append(f"轻装异常: {fast_err!r}")
            return _err(" ".join(bits))

        if not _stealth_sidecar_healthcheck(base, requests):
            return _err(_SIDECAR_UNREACHABLE_MSG)

        try:
            r = requests.post(
                endpoint,
                json={"url": url},
                timeout=http_timeout,
            )
        except Exception as e:
            try:
                import requests as req

                if isinstance(e, req.exceptions.RequestException):
                    return _err(_SIDECAR_UNREACHABLE_MSG)
            except ImportError:
                pass
            return _err(str(e))

        if not r.ok:
            detail = ""
            try:
                j = r.json()
                detail = str(j.get("detail", j))
            except Exception:
                detail = (r.text or "")[:800]
            return _err(f"重装旁路返回 HTTP {r.status_code}: {detail}")

        try:
            data = r.json()
        except Exception:
            return _err("重装旁路返回非 JSON，无法解析")

        if not isinstance(data, dict):
            return _err("重装旁路返回格式异常")

        text = data.get("text", "")
        html_excerpt = data.get("html_excerpt", "")
        return _ok(
            {
                "url": url,
                "content": {
                    "text": text,
                    "html_excerpt": html_excerpt,
                    "http_status": data.get("http_status"),
                    "via": "sidecar_heavy",
                    "in_process_error": repr(fast_err) if fast_err else None,
                },
            }
        )
    except Exception as e:
        return _err(e)


def run_dns_lookup(**kwargs: Any) -> dict[str, Any]:
    """util:dns_lookup — socket.gethostbyname_ex"""
    try:
        domain = str(kwargs.get("domain") or "").strip()
        if not domain:
            return _err("domain 不能为空")
        hostname, aliaslist, ipaddrlist = socket.gethostbyname_ex(domain)
        return _ok({"domain": domain, "hostname": hostname, "aliases": list(aliaslist), "ips": list(ipaddrlist)})
    except Exception as e:
        return _err(e)


# wttr.in 对部分中文城市名会 500；优先尝试英文名。键：用户常用中文/拼音/马来地区。
_WTTR_CITY_ALIASES: dict[str, str] = {
    "杭州": "Hangzhou",
    "北京": "Beijing",
    "上海": "Shanghai",
    "广州": "Guangzhou",
    "深圳": "Shenzhen",
    "成都": "Chengdu",
    "南京": "Nanjing",
    "武汉": "Wuhan",
    "西安": "Xian",
    "苏州": "Suzhou",
    "重庆": "Chongqing",
    "天津": "Tianjin",
    # 马来西亚（含简体/繁体常用写法；国家名默认首都）
    "马来西亚": "Kuala Lumpur",
    "大马": "Kuala Lumpur",
    "吉隆坡": "Kuala Lumpur",
    "槟城": "Penang",
    "檳城": "Penang",
    "乔治市": "George Town",
    "喬治市": "George Town",
    "新山": "Johor Bahru",
    "柔佛新山": "Johor Bahru",
    "马六甲": "Malacca",
    "麻六甲": "Malacca",
    "古晋": "Kuching",
    "古晉": "Kuching",
    "亚庇": "Kota Kinabalu",
    "亞庇": "Kota Kinabalu",
    "哥打京那巴鲁": "Kota Kinabalu",
    "哥打京那巴魯": "Kota Kinabalu",
    "布城": "Putrajaya",
    "纳闽": "Labuan",
    "納閩": "Labuan",
    "怡保": "Ipoh",
    "关丹": "Kuantan",
    "關丹": "Kuantan",
    "美里": "Miri",
    "芙蓉": "Seremban",
    "八打灵再也": "Petaling Jaya",
    "八打靈再也": "Petaling Jaya",
    "莎阿南": "Shah Alam",
    "梳邦": "Subang Jaya",
}


def _open_meteo_country_hint(city: str, candidates: list[str]) -> str | None:
    """若查询指向马来西亚，则地理编码加 countryCode=MY，避免与其它国家同名地点混淆。"""
    parts = [city] + list(candidates)
    blob = " ".join(parts)
    blob_l = blob.lower()
    my_cn = (
        "马来西亚",
        "大马",
        "吉隆坡",
        "槟城",
        "檳城",
        "乔治市",
        "喬治市",
        "新山",
        "柔佛",
        "马六甲",
        "麻六甲",
        "古晋",
        "古晉",
        "亚庇",
        "亞庇",
        "哥打京那巴鲁",
        "哥打京那巴魯",
        "布城",
        "纳闽",
        "納閩",
        "怡保",
        "关丹",
        "關丹",
        "美里",
        "芙蓉",
        "八打灵再也",
        "八打靈再也",
        "莎阿南",
        "梳邦",
        "砂拉越",
        "沙巴",
        "登嘉楼",
        "登嘉樓",
        "霹雳",
        "霹靂",
        "雪兰莪",
        "雪蘭莪",
        "彭亨",
        "吉打",
        "森美兰",
        "森美蘭",
        "玻璃市",
        "吉兰丹",
        "吉蘭丹",
        "兰卡威",
        "蘭卡威",
    )
    if any(x in blob for x in my_cn):
        return "MY"
    my_en = (
        "kuala lumpur",
        "malaysia",
        "johor bahru",
        "penang",
        "george town",
        "melaka",
        "malacca",
        "kuching",
        "kota kinabalu",
        "putrajaya",
        "labuan",
        "ipoh",
        "seremban",
        "kuantan",
        "miri",
        "petaling jaya",
        "shah alam",
        "subang jaya",
        "langkawi",
        "sarawak",
        "sabah",
        "terengganu",
        "perak",
        "selangor",
        "pahang",
        "kedah",
        "negeri sembilan",
        "perlis",
        "kelantan",
    )
    if any(w in blob_l for w in my_en):
        return "MY"
    return None


def _wttr_first_condition_dict(current_condition: Any) -> dict[str, Any]:
    """wttr j1 的 current_condition 可能为 [null] 或含非 dict，避免 NoneType.get。"""
    if not isinstance(current_condition, list):
        return {}
    for item in current_condition:
        if isinstance(item, dict):
            return item
    return {}


def _wttr_weather_desc_value(cur: dict[str, Any]) -> Any:
    wd = cur.get("weatherDesc")
    if not isinstance(wd, list) or not wd:
        return None
    first = wd[0]
    return first.get("value") if isinstance(first, dict) else None


def _wttr_nearest_area_name(data: dict[str, Any]) -> str:
    area = data.get("nearest_area")
    if not isinstance(area, list) or not area:
        return ""
    a0 = area[0]
    if not isinstance(a0, dict):
        return ""
    n = a0.get("areaName") or []
    if not isinstance(n, list) or not n:
        return ""
    z = n[0]
    return str(z.get("value", "")) if isinstance(z, dict) else ""


# Open-Meteo WMO weathercode（简版中文，与 api 文档一致）
_OM_WMO_ZH: dict[int, str] = {
    0: "晴",
    1: "大部晴朗",
    2: "少云",
    3: "阴",
    45: "雾",
    48: "雾凇雾",
    51: "小毛毛雨",
    53: "中毛毛雨",
    55: "大毛毛雨",
    61: "小雨",
    63: "中雨",
    65: "大雨",
    71: "小雪",
    73: "中雪",
    75: "大雪",
    80: "阵雨",
    81: "强阵雨",
    95: "雷暴",
    96: "雷暴伴冰雹",
    99: "强雷暴伴冰雹",
}


def _open_meteo_wmo_zh(code: Any) -> str | None:
    try:
        c = int(code)  # type: ignore[arg-type]
    except (TypeError, ValueError):
        return None
    return _OM_WMO_ZH.get(c) or f"天气代码{c}"


def _fetch_open_meteo_lite(geo_query: str, *, country_code: str | None = None) -> dict[str, Any] | None:
    """免 Key：地理编码 + 当前实况。失败返回 None。country_code 如 MY 可提高马来西亚命中率。"""
    extra = f"&countryCode={country_code}" if country_code else ""
    # 带国家码时用英文检索更稳；否则保留中文利于国内地名
    lang = "en" if country_code else "zh"
    gurl = (
        "https://geocoding-api.open-meteo.com/v1/search?"
        f"name={quote(geo_query)}&count=5&language={lang}&format=json{extra}"
    )
    try:
        req = Request(gurl, headers={"User-Agent": "Jachin-util-weather/1.1"})
        with urlopen(req, timeout=8) as resp:
            gdata = json.loads(resp.read().decode("utf-8", errors="replace"))
    except (HTTPError, URLError, json.JSONDecodeError, OSError, ValueError):
        return None
    if not isinstance(gdata, dict):
        return None
    results = gdata.get("results")
    if not isinstance(results, list) or not results:
        return None
    hit = results[0]
    if not isinstance(hit, dict):
        return None
    lat, lon = hit.get("latitude"), hit.get("longitude")
    if lat is None or lon is None:
        return None
    area_bits = [hit.get("name"), hit.get("admin1"), hit.get("country")]
    area_name = " · ".join(str(x) for x in area_bits if x)
    # 注意：current= 不能包含非法变量名；time 由响应里的 current.time 带回
    furl = (
        "https://api.open-meteo.com/v1/forecast?"
        f"latitude={lat}&longitude={lon}"
        "&current=temperature_2m,relative_humidity_2m,apparent_temperature,weather_code"
        "&timezone=auto&forecast_days=1"
    )
    try:
        req = Request(furl, headers={"User-Agent": "Jachin-util-weather/1.1"})
        with urlopen(req, timeout=8) as resp:
            fdata = json.loads(resp.read().decode("utf-8", errors="replace"))
    except (HTTPError, URLError, json.JSONDecodeError, OSError, ValueError):
        return None
    if not isinstance(fdata, dict):
        return None
    cur = fdata.get("current")
    if not isinstance(cur, dict):
        return None
    t = cur.get("temperature_2m")
    ap = cur.get("apparent_temperature")
    rh = cur.get("relative_humidity_2m")
    wc = cur.get("weather_code")
    tm = cur.get("time")
    return {
        "queried_as": geo_query,
        "area_name": area_name,
        "temp_C": None if t is None else str(round(float(t), 1)),
        "FeelsLikeC": None if ap is None else str(round(float(ap), 1)),
        "weatherDesc": _open_meteo_wmo_zh(wc),
        "humidity": None if rh is None else str(int(rh)),
        "observation_time": str(tm) if tm else None,
        "source": "open-meteo",
    }


def run_get_weather_lite(**kwargs: Any) -> dict[str, Any]:
    """util:get_weather_lite — 优先 wttr.in；失败则 Open-Meteo（均免 Key）。"""
    try:
        raw_city = kwargs.get("city") or kwargs.get("location")
        city = str(raw_city or "Beijing").strip() or "Beijing"
        candidates: list[str] = [city]
        alias = _WTTR_CITY_ALIASES.get(city)
        if alias and alias not in candidates:
            candidates.append(alias)

        last_err: str | None = None
        for q in candidates:
            enc = quote(q, safe="")
            url = f"https://wttr.in/{enc}?format=j1&lang=en"
            req = Request(url, headers={"User-Agent": "Jachin-util-weather/1.1"})
            try:
                with urlopen(req, timeout=8) as resp:
                    raw = resp.read().decode("utf-8", errors="replace")
            except HTTPError as he:
                last_err = f"HTTP {he.code}: {he.reason or 'Error'}"
                continue
            except URLError as e:
                last_err = f"网络错误: {e}"
                break
            except Exception as e:
                last_err = str(e)
                continue
            try:
                data = json.loads(raw)
            except json.JSONDecodeError as e:
                last_err = f"JSON 解析失败: {e}"
                continue
            if not isinstance(data, dict):
                last_err = "天气接口返回非 JSON 对象"
                continue
            cur = _wttr_first_condition_dict(data.get("current_condition"))
            if not cur:
                last_err = "current_condition 无有效数据（接口可能异常）"
                continue
            return _ok(
                {
                    "city_query": city,
                    "queried_as": q,
                    "area_name": _wttr_nearest_area_name(data),
                    "temp_C": cur.get("temp_C"),
                    "FeelsLikeC": cur.get("FeelsLikeC"),
                    "weatherDesc": _wttr_weather_desc_value(cur),
                    "humidity": cur.get("humidity"),
                    "observation_time": cur.get("observation_time"),
                    "source": "wttr.in",
                }
            )
        # Open-Meteo 对「杭州」等中文会先命中同名小地名；有英文别名时优先用英文名做地理编码
        om_order: list[str] = []
        if alias:
            om_order.append(alias)
        for cq in candidates:
            if cq not in om_order:
                om_order.append(cq)
        cc_hint = _open_meteo_country_hint(city, candidates)
        for cq in om_order:
            om = _fetch_open_meteo_lite(cq, country_code=cc_hint)
            if om:
                om["city_query"] = city
                return _ok(om)
        return _err(last_err or "天气查询失败（主源与备用源均不可用）")
    except Exception as e:
        return _err(e)


# ---------------------------------------------------------------------------
# 类别四：产品 / 策划 — A/B、脏数据、文本 diff、漏斗 ROI
# ---------------------------------------------------------------------------


def run_ab_test_calc(**kwargs: Any) -> dict[str, Any]:
    """
    util:ab_test_calc — 两组转化率 Z 检验（合并方差），双尾 P-value；α=0.05 时报告 is_significant。
    JSON：visitors_a, conversions_a, visitors_b, conversions_b（均为 int）
    """
    try:
        va = int(kwargs.get("visitors_a", 0))
        ca = int(kwargs.get("conversions_a", 0))
        vb = int(kwargs.get("visitors_b", 0))
        cb = int(kwargs.get("conversions_b", 0))
        if va <= 0 or vb <= 0:
            return _err("visitors_a / visitors_b 必须为正整数")
        if ca < 0 or cb < 0 or ca > va or cb > vb:
            return _err("conversions 必须在 [0, visitors] 范围内")
        p1 = ca / va
        p2 = cb / vb
        p_pool = (ca + cb) / (va + vb)
        if p_pool <= 0 or p_pool >= 1:
            # 全 0 或全 1 时方差为 0
            z = 0.0
            p_val = 1.0
        else:
            se_sq = p_pool * (1.0 - p_pool) * (1.0 / va + 1.0 / vb)
            if se_sq <= 0:
                return _err("无法计算标准误（请检查样本量）")
            se = math.sqrt(se_sq)
            z = (p1 - p2) / se
            # 标准正态双尾：P(|Z|>|z|) = erfc(|z|/√2)
            p_val = float(math.erfc(abs(z) / math.sqrt(2.0)))
        alpha = 0.05
        is_sig = bool(p_val < alpha)
        return _ok(
            {
                "conversion_rate_a": round(p1, 6),
                "conversion_rate_b": round(p2, 6),
                "z_score": round(z, 6),
                "p_value_two_tailed": round(p_val, 8),
                "confidence_level": 0.95,
                "alpha": alpha,
                "is_significant": is_sig,
            }
        )
    except Exception as e:
        return _err(e)


def run_fake_data_gen(**kwargs: Any) -> dict[str, Any]:
    """util:fake_data_gen — Faker 生成占位数据。JSON：locale（默认 zh_CN）, count（默认 5，最大 50）, fields（name|phone|email|address|company|job）"""
    try:
        try:
            from faker import Faker
        except ImportError:
            return _err("未安装 Faker，请执行: pip install faker")

        locale = str(kwargs.get("locale") or "zh_CN").strip() or "zh_CN"
        count = int(kwargs.get("count", 5))
        count = max(1, min(50, count))
        raw_fields = kwargs.get("fields")
        allowed = {"name", "phone", "email", "address", "company", "job"}
        if raw_fields is None or (isinstance(raw_fields, list) and len(raw_fields) == 0):
            use_fields = sorted(allowed)
        elif isinstance(raw_fields, list):
            use_fields = []
            for x in raw_fields:
                k = str(x).strip()
                if k not in allowed:
                    return _err(f"未知字段: {k!r}，允许: {sorted(allowed)}")
                if k not in use_fields:
                    use_fields.append(k)
        else:
            return _err("fields 须为字符串数组")

        fake = Faker(locale)
        gen_map = {
            "name": lambda: fake.name(),
            "phone": lambda: fake.phone_number(),
            "email": lambda: fake.email(),
            "address": lambda: fake.address().replace("\n", ", "),
            "company": lambda: fake.company(),
            "job": lambda: fake.job(),
        }
        dummy_data: list[dict[str, str]] = []
        for _ in range(count):
            row: dict[str, str] = {}
            for f in use_fields:
                row[f] = str(gen_map[f]())
            dummy_data.append(row)
        return _ok({"locale": locale, "count": count, "fields": use_fields, "dummy_data": dummy_data})
    except Exception as e:
        return _err(e)


def run_text_diff(**kwargs: Any) -> dict[str, Any]:
    """util:text_diff — difflib.unified_diff，仅返回以 + / - 开头的实质差异行（去掉文件头）。"""
    try:
        text1 = str(kwargs.get("text1", ""))
        text2 = str(kwargs.get("text2", ""))
        a = text1.splitlines()
        b = text2.splitlines()
        diff_iter = difflib.unified_diff(a, b, lineterm="")
        change_lines: list[str] = []
        for line in diff_iter:
            if not line:
                continue
            if line.startswith("---") or line.startswith("+++"):
                continue
            if line.startswith("@@"):
                continue
            if line.startswith("+") or line.startswith("-"):
                change_lines.append(line)
        return _ok(
            {
                "diff_lines": change_lines,
                "added_count": sum(1 for x in change_lines if x.startswith("+")),
                "removed_count": sum(1 for x in change_lines if x.startswith("-")),
            }
        )
    except Exception as e:
        return _err(e)


def run_funnel_calc(**kwargs: Any) -> dict[str, Any]:
    """
    util:funnel_calc — 链式漏斗人数；可选 CAC（单访客获客成本）与 ARPU（末层单用户收入）算 ROI。
    JSON：initial_traffic, conversion_rates（每层 0~1）；可选 cac, arpu
    """
    try:
        n0 = int(kwargs.get("initial_traffic", 0))
        if n0 <= 0:
            return _err("initial_traffic 必须为正整数")
        rates_raw = kwargs.get("conversion_rates")
        if not isinstance(rates_raw, list) or len(rates_raw) == 0:
            return _err("conversion_rates 须为非空数组")
        rates: list[float] = []
        for i, x in enumerate(rates_raw):
            r = float(x)
            if r < 0 or r > 1:
                return _err(f"conversion_rates[{i}] 应在 [0,1] 内")
            rates.append(r)

        layer_counts: list[float] = [float(n0)]
        cur = float(n0)
        for r in rates:
            cur = cur * r
            layer_counts.append(cur)

        final_users = layer_counts[-1]
        out: dict[str, Any] = {
            "layer_counts": [round(x, 6) for x in layer_counts],
            "final_conversions": round(final_users, 6),
        }

        cac_raw = kwargs.get("cac")
        arpu_raw = kwargs.get("arpu")
        if cac_raw is None and arpu_raw is None:
            out["cpa"] = None
            out["total_cost"] = None
            out["total_revenue"] = None
            out["roi"] = None
            return _ok(out)

        if cac_raw is None or arpu_raw is None:
            return _err("计算财务指标时请同时提供 cac 与 arpu（均为数字）")

        cac = float(cac_raw)
        arpu = float(arpu_raw)
        if cac < 0 or arpu < 0:
            return _err("cac / arpu 不可为负")

        total_cost = cac * float(n0)
        total_revenue = arpu * final_users
        out["total_cost"] = round(total_cost, 6)
        out["total_revenue"] = round(total_revenue, 6)
        if final_users > 0:
            out["cpa"] = round(total_cost / final_users, 6)
        else:
            out["cpa"] = None
        if total_cost > 0:
            out["roi"] = round((total_revenue - total_cost) / total_cost, 6)
        else:
            out["roi"] = None
        return _ok(out)
    except Exception as e:
        return _err(e)


# ---------------------------------------------------------------------------
# 类别五：本机可见提醒（与 com.jachin.os-mate.desktop_notify 能力对齐）
# ---------------------------------------------------------------------------


def run_desktop_message_box(**kwargs: Any) -> dict[str, Any]:
    """
    util:desktop_message_box — 在本机弹出**立即**可见的提醒。

    - Windows: WinForms MessageBox（异步 Popen，不阻塞 L3）。
    - macOS: osascript display notification（非阻塞）。
    - Linux: notify-send（若存在）。

    标题/正文经 Base64 传入 PowerShell，避免引号注入。
    「到 18:10 再弹」需系统计划任务 / 闹钟或到时仍运行的调度器触发本工具，本函数不负责定时。
    """
    title = (str(kwargs.get("title") or "Jachin").strip()[:200] or "Jachin").strip()
    message = str(kwargs.get("message") or "").strip()
    if not message:
        return _err("message 不能为空")
    message = message[:4000]

    system = platform.system()
    try:
        if system == "Windows":
            tb = base64.b64encode(title.encode("utf-8")).decode("ascii")
            mb = base64.b64encode(message.encode("utf-8")).decode("ascii")
            ps = (
                "Add-Type -AssemblyName System.Windows.Forms; "
                f"$t=[System.Text.Encoding]::UTF8.GetString([System.Convert]::FromBase64String('{tb}')); "
                f"$m=[System.Text.Encoding]::UTF8.GetString([System.Convert]::FromBase64String('{mb}')); "
                "[void][System.Windows.Forms.MessageBox]::Show($m,$t)"
            )
            cf = getattr(subprocess, "CREATE_NO_WINDOW", 0)
            subprocess.Popen(
                ["powershell", "-NoProfile", "-WindowStyle", "Hidden", "-Command", ps],
                creationflags=cf,
            )
            return _ok(
                "已异步弹出 Windows 消息框（用户点击确定后关闭）。"
                "若用户要求「指定时刻」提醒，须说明本工具仅立即弹出，到时需闹钟/计划任务或调度触发。"
            )
        if system == "Darwin":
            safe_t = title.replace("\\", "\\\\").replace('"', '\\"')[:120]
            safe_m = message.replace("\\", "\\\\").replace('"', '\\"')[:500]
            subprocess.Popen(
                ["osascript", "-e", f'display notification "{safe_m}" with title "{safe_t}"']
            )
            return _ok("已发送 macOS 通知（非模态）。")
        subprocess.Popen(
            ["notify-send", "-a", "Jachin", title[:128], message[:512]],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        return _ok("已调用 notify-send（无图形会话时可能不可见）。")
    except FileNotFoundError:
        return _err("当前环境未找到 notify-send / osascript，或 PowerShell 不可用")
    except Exception as e:
        return _err(str(e))


def run_schedule_desktop_reminder(**kwargs: Any) -> dict[str, Any]:
    """
    util:schedule_desktop_reminder — 向本机 Jachin 桌面客户端注册**定时**右下角哨兵提醒（HTTP 127.0.0.1:8002）。

    须 Jachin 桌面已运行并监听该端口。与 util:desktop_message_box（立即弹窗）不同。
    时刻三选一：fire_at_unix_ms | delay_seconds | fire_at_iso（ISO8601，缺省按 Asia/Shanghai 解释无时区的本地时刻）。
    """
    title = (str(kwargs.get("title") or "Jachin").strip()[:200] or "Jachin").strip()
    body = str(kwargs.get("body") or kwargs.get("message") or "").strip()
    if not body:
        return _err("body 或 message 不能为空")
    body = body[:4000]

    fire_at_unix_ms = kwargs.get("fire_at_unix_ms")
    delay_seconds = kwargs.get("delay_seconds")
    fire_at_iso = kwargs.get("fire_at_iso")
    filled = 0
    if fire_at_unix_ms is not None and str(fire_at_unix_ms).strip() != "":
        filled += 1
    if delay_seconds is not None and str(delay_seconds).strip() != "":
        filled += 1
    if fire_at_iso is not None and str(fire_at_iso).strip() != "":
        filled += 1
    if filled != 1:
        return _err("fire_at_unix_ms、delay_seconds、fire_at_iso 须且只能指定其一")

    now_ms = int(time.time() * 1000)
    target_ms: int
    try:
        if fire_at_unix_ms is not None and str(fire_at_unix_ms).strip() != "":
            target_ms = int(fire_at_unix_ms)
        elif delay_seconds is not None and str(delay_seconds).strip() != "":
            target_ms = now_ms + int(float(delay_seconds) * 1000)
        else:
            s = str(fire_at_iso).strip()
            if s.endswith("Z"):
                s = s[:-1] + "+00:00"
            dt = datetime.fromisoformat(s)
            try:
                from zoneinfo import ZoneInfo

                sh = ZoneInfo("Asia/Shanghai")
            except Exception:
                sh = timezone(timedelta(hours=8))
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=sh)
            else:
                dt = dt.astimezone(sh)
            target_ms = int(dt.timestamp() * 1000)
    except (ValueError, TypeError, OSError) as e:
        return _err(f"时刻解析失败: {e}")

    base_url = (
        os.environ.get("JACHIN_DESKTOP_REMINDER_URL") or "http://127.0.0.1:8002/jachin/v1/reminders"
    ).rstrip("/")
    payload = json.dumps(
        {"fire_at_unix_ms": target_ms, "title": title, "body": body},
        ensure_ascii=False,
    ).encode("utf-8")
    req = Request(
        base_url,
        data=payload,
        method="POST",
        headers={"Content-Type": "application/json; charset=utf-8"},
    )
    try:
        with urlopen(req, timeout=8.0) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
    except HTTPError as e:
        try:
            err_body = e.read().decode("utf-8", errors="replace")
        except Exception:
            err_body = ""
        return _err(f"HTTP {e.code} {e.reason} {err_body}".strip())
    except URLError as e:
        return _err(
            "无法连接 Jachin 桌面提醒接口（请确认桌面端已启动且监听 127.0.0.1:8002）。"
            f" 详情: {e.reason!s}"
        )
    except Exception as e:
        return _err(str(e))

    try:
        data = json.loads(raw) if raw.strip() else {}
    except json.JSONDecodeError:
        return _err(f"响应非 JSON: {raw[:500]}")

    if not data.get("ok"):
        return _err(str(data.get("error") or data))
    rid = data.get("id")
    return _ok(
        {
            "id": rid,
            "fire_at_unix_ms": target_ms,
            "hint": "提醒已写入桌面端；到点将弹出右下角哨兵（与立即弹窗的 util:desktop_message_box 不同）。",
        }
    )


def run_lark_send_text(**kwargs: Any) -> dict[str, Any]:
    """
    util:lark_send_text — 经 Lark Open API 向指定会话发送纯文本（**通用** LARK_APP_ID / LARK_APP_SECRET；
    与 HR 多维表分离时请把招聘应用填到 HR_LARK_APP_*，勿与终端用户 open_id 混用）。

    典型链路：用户要「总结某网页并发到飞书」→ 先用 **util:stealth_extract**（或 MCP）取正文，再在模型内压缩后调用本工具 **text**。
    **接收者** 解析顺序（与飞书「发送消息」一致）：
    1) Action Input 里同时给出 **receive_id_type**（非默认 chat_id）与 **chat_id/receive_id** 时，按模型指定发；
    2) 环境变量 **LARK_USER_OPEN_ID** / **LARK_DM_OPEN_ID**（用户 **ou_** open_id）— 官方推荐用于**机器人私聊**，不必依赖 oc_ 会话 ID；
    3) **LARK_CHAT_ID**（群或会话 oc_…）；
    4) Action 仅含正文时的 **chat_id**；
    5) 当前轮 ContextVar（终端镜像）。
    """
    try:
        from l3_node.channels.lark.client import _ensure_dotenv_loaded

        _ensure_dotenv_loaded()
    except Exception:
        pass

    text = str(kwargs.get("text") or kwargs.get("message") or "").strip()
    if not text:
        return _err("text 不能为空（可将摘要作为 text 传入）")
    # 飞书单条文本上限约 6k，留余量
    text = text[:5900]

    _valid_recv = ("chat_id", "open_id", "user_id", "union_id", "email")
    _kw_cid = str(kwargs.get("chat_id") or kwargs.get("receive_id") or "").strip()
    _kw_type = str(kwargs.get("receive_id_type") or "").strip().lower()
    _p2p_open = (
        os.environ.get("LARK_USER_OPEN_ID")
        or os.environ.get("LARK_DM_OPEN_ID")
        or ""
    ).strip()
    _env_cid = (
        os.environ.get("LARK_CHAT_ID")
        or os.environ.get("LARK_DEFAULT_CHAT_ID")
        or os.environ.get("FEISHU_CHAT_ID")
        or ""
    ).strip()

    chat_id = ""
    receive_id_type = "chat_id"
    if _kw_type and _kw_type in _valid_recv and _kw_type != "chat_id" and _kw_cid:
        chat_id = _kw_cid
        receive_id_type = _kw_type
    elif _p2p_open:
        chat_id = _p2p_open
        receive_id_type = "open_id"
    elif _env_cid:
        chat_id = _env_cid
        receive_id_type = "chat_id"
    elif _kw_cid:
        chat_id = _kw_cid
        receive_id_type = _kw_type if _kw_type in _valid_recv else "chat_id"
    else:
        try:
            from l3_node.channels.lark.turn_chat_context import peek_lark_chat_id_for_tools

            chat_id = peek_lark_chat_id_for_tools()
        except Exception:
            chat_id = ""
        receive_id_type = "chat_id"
    if not chat_id:
        return _err(
            "未指定接收者：请配置 LARK_CHAT_ID（群/会话 oc_…），或私聊配置 LARK_USER_OPEN_ID（用户 ou_…），"
            "或在 Action Input 传入 chat_id / receive_id_type。"
        )

    if receive_id_type not in _valid_recv:
        receive_id_type = "chat_id"

    try:
        from l3_node.channels.lark.client import get_lark_api_base, get_tenant_access_token, resolve_lark_credentials
        from l3_node.channels.lark.im import send_text as lark_send_text_impl
        from l3_node.channels.lark.receive_resolve import normalize_lark_im_receive
    except ImportError as e:
        return _err(f"L3 Lark 通道未就绪: {e}")

    aid, sec, yb = resolve_lark_credentials()
    if not aid or not sec:
        return _err(
            "未配置 LARK_APP_ID / LARK_APP_SECRET（或 ~/.jachin/config/im_channels.yaml 中 im_channels.lark）。"
            "参见仓库根 .env.example。"
        )
    base = yb or get_lark_api_base()
    try:
        tkn = get_tenant_access_token(app_id=aid, app_secret=sec, api_base=base)
        n_rid, n_rt, n_err = normalize_lark_im_receive(chat_id, receive_id_type, token=tkn, api_base=base)
        if n_err:
            return _err(n_err)
        chat_id = n_rid
        receive_id_type = n_rt
    except Exception as e:
        return _err(str(e))
    res = lark_send_text_impl(
        chat_id,
        text,
        receive_id_type=receive_id_type,
        app_id=aid,
        app_secret=sec,
        api_base=base,
    )
    if res.get("status") == "success":
        return _ok(
            {
                "lark": res.get("msg", "已送达"),
                "receive_id_type": receive_id_type,
            }
        )
    err_raw = str(res.get("error") or res)
    _el = err_raw.lower()
    if "cross app" in _el:
        return _err(
            f"{err_raw}。"
            f" 当前 tenant 对应应用 ID：{aid}；接收方 open_id 须与该应用在飞书开放平台为同一应用下解析得到。"
            " 若已在 skills_repo/plugin/.env 配置通用机器人仍报错，请检查系统/用户环境变量是否残留其它 LARK_APP_ID，"
            "或确认 ou_ 不是从其它应用复制；可用 JACHIN_IGNORE_PLUGIN_LARK=1 禁止 plugin/.env 覆盖环境变量。"
        )
    if "out of the chat" in _el or "no availability" in _el or "230013" in err_raw:
        return _err(
            f"{err_raw}。"
            "常见原因：(1) 群聊：机器人未入群，请 @ 添加该应用机器人；"
            "(2) 私聊：飞书文档推荐用用户 **open_id**（ou_…）发消息，可在 .env 设置 LARK_USER_OPEN_ID=你的 ou_…，"
            "或在 Action 中传 receive_id_type=open_id、receive_id=ou_…，避免仅用 oc_ 会话 ID；"
            "(3) 用户在应用「可用范围」外（开发者后台调整可见范围并重新发布）；"
            "(4) 终端镜像 oc_ 与目标不一致时，用 LARK_CHAT_ID 或 LARK_USER_OPEN_ID 固定目标。"
        )
    return _err(err_raw)


def run_lark_resolve_user(**kwargs: Any) -> dict[str, Any]:
    """
    util:lark_resolve_user — 通过邮箱或手机号解析飞书用户 **open_id**（contact/v3/users/batch_get_id）。
    需应用开通通讯录相关权限；仅有姓名时优先 **util:lark_search_user**，或别名表 / 邮箱 / ou_。
    """
    try:
        from l3_node.channels.lark.client import _ensure_dotenv_loaded

        _ensure_dotenv_loaded()
    except Exception:
        pass
    try:
        from l3_node.channels.lark.client import get_lark_api_base, get_tenant_access_token, resolve_lark_credentials
        from l3_node.channels.lark.receive_resolve import (
            feishu_batch_get_id,
            is_contact_scope_denied_payload,
            looks_like_email,
            normalize_mobile_for_feishu,
            pick_open_id_from_batch_get,
            resolve_display_name_to_id,
            search_user_candidates_by_name,
        )
    except ImportError as e:
        return _err(str(e))

    display_name = str(kwargs.get("display_name") or kwargs.get("name") or "").strip()
    email = str(kwargs.get("email") or "").strip()
    mobile = str(kwargs.get("mobile") or "").strip()
    if display_name and not email and not mobile:
        mapped = resolve_display_name_to_id(display_name)
        if mapped:
            if mapped.startswith("ou_"):
                return _ok(
                    {
                        "open_id": mapped,
                        "user_id": None,
                        "union_id": None,
                        "source": "LARK_DISPLAY_NAME_MAP",
                        "hint": "来自别名表；发消息：util:lark_send_text + receive_id_type=open_id",
                    }
                )
            if looks_like_email(mapped):
                email = mapped
            elif normalize_mobile_for_feishu(mapped):
                mobile = mapped
        else:
            return _err(
                "未配置该英文名的映射。请在环境变量 LARK_DISPLAY_NAME_MAP 或 ~/.jachin/lark_display_name_map.json "
                "中设置 {\"vivian\":\"ou_xxx\"}，或提供邮箱/手机。"
            )
    if not email and not mobile:
        return _err("请至少提供 email 或 mobile，或 display_name（须先在 LARK_DISPLAY_NAME_MAP 配置映射）")
    if email and not looks_like_email(email):
        return _err("email 格式不合法")
    mob_norm: str | None = None
    if mobile:
        mob_norm = normalize_mobile_for_feishu(mobile)
        if not mob_norm:
            return _err("mobile 无法解析为中国大陆 11 位手机号（可含空格/短横线）")

    aid, sec, yb = resolve_lark_credentials()
    if not aid or not sec:
        return _err("未配置 LARK_APP_ID / LARK_APP_SECRET（或 im_channels.yaml）。")
    base = yb or get_lark_api_base()
    try:
        tkn = get_tenant_access_token(app_id=aid, app_secret=sec, api_base=base)
    except Exception as e:
        return _err(str(e))

    data = feishu_batch_get_id(
        token=tkn,
        api_base=base,
        emails=[email] if email else None,
        mobiles=[mob_norm] if mob_norm else None,
    )
    oid, err = pick_open_id_from_batch_get(data)
    _msg_l = str(data.get("msg") or "").lower()
    _need_token_retry = (err or not oid) and (
        is_contact_scope_denied_payload(data)
        or data.get("code") == 403
        or ("access denied" in _msg_l and "scope" in _msg_l)
    )
    if _need_token_retry:
        import time

        try:
            from l3_node.channels.lark.client import invalidate_lark_tenant_token_cache

            invalidate_lark_tenant_token_cache()
        except Exception:
            pass
        time.sleep(2.0)
        try:
            tkn = get_tenant_access_token(app_id=aid, app_secret=sec, api_base=base)
        except Exception:
            pass
        data = feishu_batch_get_id(
            token=tkn,
            api_base=base,
            emails=[email] if email else None,
            mobiles=[mob_norm] if mob_norm else None,
        )
        oid, err = pick_open_id_from_batch_get(data)
    if err or not oid:
        # batch_get_id 报缺 contact:user.id:readonly 时，仍可能具备通讯录搜索/列表类权限：用邮箱前缀做一次降级解析
        try:
            if email and looks_like_email(email) and data and is_contact_scope_denied_payload(data):
                q = email.split("@", 1)[0].strip()
                if q:
                    cands, _serr = search_user_candidates_by_name(q, tenant_token=tkn, api_base=base)
                    if len(cands) == 1:
                        oid_fb = str((cands[0] or {}).get("open_id") or "").strip()
                        if oid_fb.startswith("ou_"):
                            return _ok(
                                {
                                    "open_id": oid_fb,
                                    "user_id": None,
                                    "union_id": None,
                                    "source": "name_search_fallback",
                                    "hint": (
                                        "batch_get_id 因 scope 被拒，已通过姓名关键词解析到唯一用户；"
                                        "发消息：util:lark_send_text + receive_id_type=open_id。"
                                        "建议在后台开通「通过手机号或邮箱获取用户 ID」并发布版本，避免依赖降级。"
                                    ),
                                }
                            )
                    if len(cands) > 1:
                        lines = [
                            f"- {u.get('name', '-') or '-'} / {u.get('en_name', '-') or '-'} | 邮箱:{u.get('email', '-') or '-'} | open_id:{u.get('open_id', '')}"
                            for u in cands[:15]
                        ]
                        return _err(
                            "batch_get_id 无 contact:user.id:readonly 权限；按邮箱前缀搜索到多名用户，请确认目标：\n"
                            + "\n".join(lines)
                        )
        except Exception:
            pass
        return _err(err or "batch_get_id 失败")
    ulist = (data.get("data") or {}).get("user_list") or []
    u0 = ulist[0] if ulist else {}
    return _ok(
        {
            "open_id": oid,
            "user_id": (u0.get("user_id") or "").strip() or None,
            "union_id": (u0.get("union_id") or "").strip() or None,
            "hint": "发消息：util:lark_send_text 传入 receive_id_type=open_id，chat_id/receive_id 填上述 open_id",
        }
    )


def run_lark_search_user(**kwargs: Any) -> dict[str, Any]:
    """
    util:lark_search_user — 按姓名/昵称在飞书通讯录中搜索用户（search/v1/user，可选 user token；失败则 contact/v3/users 分页匹配）。
    唯一候选时可将 open_id 用于 util:lark_send_text；多个时须在回复中列出部门/邮箱并请用户确认。
    """
    q = str(kwargs.get("query") or kwargs.get("name") or "").strip()
    if not q:
        return _err("query 或 name 必填")
    try:
        from l3_node.channels.lark.client import (
            _ensure_dotenv_loaded,
            get_lark_api_base,
            get_tenant_access_token,
            resolve_lark_credentials,
        )
        from l3_node.channels.lark.receive_resolve import search_user_candidates_by_name
    except ImportError as e:
        return _err(str(e))
    try:
        _ensure_dotenv_loaded()
    except Exception:
        pass
    aid, sec, yb = resolve_lark_credentials()
    if not aid or not sec:
        return _err("未配置 LARK_APP_ID / LARK_APP_SECRET（或 im_channels.yaml）。")
    base = yb or get_lark_api_base()
    try:
        tkn = get_tenant_access_token(app_id=aid, app_secret=sec, api_base=base)
    except Exception as e:
        return _err(str(e))
    cands, serr = search_user_candidates_by_name(q, tenant_token=tkn, api_base=base)
    if not cands:
        msg = serr or f"未能在飞书中搜到名为「{q}」的用户，请确认姓名是否有误，或配置 LARK_FEISHU_USER_ACCESS_TOKEN 以启用搜索接口。"
        return _err(msg)
    return _ok({"query": q, "count": len(cands), "candidates": cands})


# ---------------------------------------------------------------------------
# 类别六：系统健康
# ---------------------------------------------------------------------------


def run_health_stats(**kwargs: Any) -> dict[str, Any]:
    """sys:health_stats — CPU / 内存 / 磁盘（需 psutil）。"""
    try:
        _ = kwargs
        try:
            import psutil
        except ImportError:
            return _err("未安装 psutil，请执行: pip install psutil")

        cpu = psutil.cpu_percent(interval=0.15)
        vm = psutil.virtual_memory()
        du = psutil.disk_usage(str(Path.home()))
        return _ok(
            {
                "cpu_percent": round(float(cpu), 2),
                "memory_free_mb": round(vm.available / (1024 * 1024), 2),
                "memory_total_mb": round(vm.total / (1024 * 1024), 2),
                "disk_free_gb": round(du.free / (1024**3), 3),
                "disk_total_gb": round(du.total / (1024**3), 3),
            }
        )
    except Exception as e:
        return _err(e)


def run_list_env_safe(**kwargs: Any) -> dict[str, Any]:
    """sys:list_env_safe — 仅环境变量名，不含值。"""
    try:
        _ = kwargs
        names = sorted(os.environ.keys())
        return _ok({"count": len(names), "keys": names})
    except Exception as e:
        return _err(e)


# ---------------------------------------------------------------------------
# Office：Word / Excel 原生二进制（util:generate_office_doc）
# ---------------------------------------------------------------------------

try:
    import docx  # noqa: F401
    from docx import Document as _DocxDocument
    from docx.shared import Pt, RGBColor  # noqa: F401

    _OFFICE_DOCX_OK = True
except ImportError:
    docx = None  # type: ignore[assignment]
    _DocxDocument = None  # type: ignore[misc]
    Pt = RGBColor = None  # type: ignore[misc]
    _OFFICE_DOCX_OK = False

try:
    import openpyxl  # noqa: F401
    from openpyxl import Workbook as _OpenpyxlWorkbook

    _OFFICE_XLSX_OK = True
except ImportError:
    openpyxl = None  # type: ignore[assignment]
    _OpenpyxlWorkbook = None  # type: ignore[misc]
    _OFFICE_XLSX_OK = False


def _resolve_safe_output_path(file_path: str) -> Path:
    """与 core:fs_write 一致：输出须在 native_write_allowlist 白名单内。"""
    from l3_node.primitives.native_write_allowlist import assert_path_allowed_for_native_write
    from l3_node.workspace_context import get_effective_workspace_root

    workspace = get_effective_workspace_root()
    raw = (file_path or "").strip().replace("\\", "/")
    fp = Path(raw).expanduser()
    if not fp.is_absolute():
        fp = (workspace / fp).resolve()
    else:
        fp = fp.resolve()
    assert_path_allowed_for_native_write(fp)
    return fp


def _xlsx_safe_sheet_name(name: str) -> str:
    import re

    base = (name or "Sheet").strip() or "Sheet"
    base = re.sub(r"[\*\:\\/\?\[\]]+", "_", base)
    return base[:31] if len(base) > 31 else base


def run_generate_office_doc(**kwargs: Any) -> dict[str, Any]:
    """
    util:generate_office_doc — 格式转换层：由 content_json 渲染原生 .docx / .xlsx。
    参数：file_format（docx|xlsx，兼容旧名 file_type）、file_path、content_json（兼容旧名 content_data）。
    docx：content_json.blocks[]，块 type 为 h1|h2|h3|p|bullet|table。
    xlsx：content_json.sheets[]，每项含 sheet_name 与 data 二维数组。
    """
    try:
        fmt = str(kwargs.get("file_format") or kwargs.get("file_type") or "").strip().lower()
        if fmt not in ("docx", "xlsx"):
            return _err('file_format 须为 "docx" 或 "xlsx"（也可用旧字段 file_type）')
        raw_fp = str(kwargs.get("file_path") or "").strip()
        if not raw_fp:
            return _err("file_path 不能为空")

        if fmt == "docx" and not _OFFICE_DOCX_OK:
            return _err("缺少依赖 python-docx，请执行: pip install python-docx")
        if fmt == "xlsx" and not _OFFICE_XLSX_OK:
            return _err("缺少依赖 openpyxl，请执行: pip install openpyxl")

        fp = _resolve_safe_output_path(raw_fp)
        want_ext = ".docx" if fmt == "docx" else ".xlsx"
        if fp.suffix.lower() != want_ext:
            return _err(f"file_path 扩展名须为 {want_ext}")

        cj = kwargs.get("content_json", kwargs.get("content_data"))
        if isinstance(cj, str):
            cj = json.loads(cj)
        if not isinstance(cj, dict):
            return _err("content_json 须为 JSON 对象（或可解析的 JSON 字符串；兼容旧字段 content_data）")

        fp.parent.mkdir(parents=True, exist_ok=True)

        if fmt == "docx":
            assert _DocxDocument is not None
            doc = _DocxDocument()
            blocks = cj.get("blocks")
            if not isinstance(blocks, list):
                return _err("docx 的 content_json 须包含 blocks 数组")
            # python-docx：level 1–3 对应 Word「标题 1」–「标题 3」
            heading_level = {"h1": 1, "h2": 2, "h3": 3}
            for i, block in enumerate(blocks):
                if not isinstance(block, dict):
                    return _err(f"blocks[{i}] 须为对象")
                bt = str(block.get("type") or "").strip().lower()
                if bt in heading_level:
                    doc.add_heading(str(block.get("text") or ""), level=heading_level[bt])
                elif bt == "p":
                    doc.add_paragraph(str(block.get("text") or ""))
                elif bt == "bullet":
                    doc.add_paragraph(str(block.get("text") or ""), style="List Bullet")
                elif bt == "table":
                    data = block.get("data") or []
                    if not isinstance(data, list) or not data:
                        continue
                    if not all(isinstance(r, list) for r in data):
                        return _err(f"blocks[{i}].data 须为二维数组")
                    n_rows = len(data)
                    n_cols = max((len(r) for r in data), default=0)
                    if n_cols < 1:
                        continue
                    tbl = doc.add_table(rows=n_rows, cols=n_cols)
                    tbl.style = "Table Grid"
                    for ri, row in enumerate(data):
                        for ci in range(n_cols):
                            val = row[ci] if ci < len(row) else ""
                            tbl.cell(ri, ci).text = "" if val is None else str(val)
                else:
                    return _err(f"不支持的 block.type: {bt!r}（支持 h1,h2,h3,p,bullet,table）")
            doc.save(str(fp))
        else:
            assert _OpenpyxlWorkbook is not None
            sheets = cj.get("sheets")
            if not isinstance(sheets, list) or not sheets:
                return _err("xlsx 的 content_json 须包含非空 sheets 数组")
            wb = _OpenpyxlWorkbook()
            for si, sh in enumerate(sheets):
                if not isinstance(sh, dict):
                    return _err(f"sheets[{si}] 须为对象")
                name = str(sh.get("sheet_name") or "").strip() or f"Sheet{si + 1}"
                data = sh.get("data")
                if not isinstance(data, list):
                    return _err(f"sheets[{si}].data 须为二维数组（行列表）")
                if si == 0:
                    ws = wb.active
                    assert ws is not None
                    ws.title = _xlsx_safe_sheet_name(name)
                else:
                    ws = wb.create_sheet(title=_xlsx_safe_sheet_name(name))
                for row in data:
                    if not isinstance(row, list):
                        return _err(f"sheets[{si}].data 中每一行须为数组")
                    ws.append(row)
            wb.save(str(fp))

        return {"ok": True, "file_path": str(fp.resolve())}
    except json.JSONDecodeError as e:
        return _err(f"content_json JSON 无效: {e}")
    except Exception as e:
        return _err(e)


# ---------------------------------------------------------------------------
# 长文 Map-Reduce：逐章 LLM 后拼装 Markdown（util:compose_long_document）
# ---------------------------------------------------------------------------


def _parse_outline_sections(raw: Any) -> tuple[list[str], str | None]:
    """解析大纲节点列表。成功返回 (sections, None)；失败返回 ([], error_msg)。"""
    if raw is None:
        return [], "outline_sections 缺失"
    if isinstance(raw, list):
        sections = [str(x).strip() for x in raw if str(x).strip()]
        return (sections, None)
    if isinstance(raw, str):
        s = raw.strip()
        if not s:
            return [], "outline_sections 为空"
        if s.startswith("["):
            try:
                parsed = json.loads(s)
            except json.JSONDecodeError as e:
                return [], f"outline_sections JSON 无效: {e}"
            if not isinstance(parsed, list):
                return [], "outline_sections 须为非空 JSON 数组"
            sections = [str(x).strip() for x in parsed if str(x).strip()]
            return (sections, None)
        sections = [ln.strip() for ln in s.replace("\r", "").split("\n") if ln.strip()]
        return (sections, None)
    return [], "outline_sections 须为字符串数组或可解析的 JSON 数组字符串"


def _run_coroutine_in_fresh_loop(coro):
    """在无非运行事件循环的上下文中执行协程；若当前线程已有 loop，则放到独立线程里 asyncio.run。"""
    import asyncio

    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(coro)

    import concurrent.futures

    def _entry() -> Any:
        return asyncio.run(coro)

    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
        fut = ex.submit(_entry)
        return fut.result()


def run_compose_long_document(**kwargs: Any) -> dict[str, Any]:
    """
    util:compose_long_document — Map-Reduce 式长文：按大纲逐章调用 LiteLLM（无状态），再拼装为 Markdown 落盘。
    参数：file_path、topic、outline_sections（字符串列表，或 JSON 数组字符串）。
    """
    try:
        topic = str(kwargs.get("topic") or "").strip()
        if not topic:
            return _err("topic 不能为空")
        raw_fp = str(kwargs.get("file_path") or "").strip()
        if not raw_fp:
            return _err("file_path 不能为空")

        sections, sec_err = _parse_outline_sections(kwargs.get("outline_sections"))
        if sec_err:
            return _err(sec_err)
        if not sections:
            return _err("outline_sections 不能为空")
        if len(sections) > 24:
            return _err("outline_sections 章节过多（建议 5–12 节，上限 24）")

        from core.llm_provider import LiteLLMEngine, get_complex_model_litellm_id

        try:
            eng = LiteLLMEngine(model_name=get_complex_model_litellm_id())
        except Exception:
            eng = LiteLLMEngine()

        parts: list[str] = []
        for section in sections:
            chapter_prompt = (
                f"你正在撰写《{topic}》，当前章节为【{section}】。"
                "请输出极其详尽的专业内容，务必展开论述，不要输出总结套话。"
                "直接返回 Markdown 文本。"
            )

            async def _one_chapter(
                _prompt: str = chapter_prompt,
            ) -> str:
                out = await eng.generate_response(
                    [{"role": "user", "content": _prompt}],
                    tools=None,
                    temperature=0.75,
                    max_tokens=8192,
                    call_purpose="util_compose_long_document",
                )
                if isinstance(out, dict):
                    return str(out.get("content") or "").strip()
                return str(out or "").strip()

            chapter = _run_coroutine_in_fresh_loop(_one_chapter())
            parts.append(f"## {section}\n\n{chapter}")

        body = "\n\n".join(parts)
        fp = _resolve_safe_output_path(raw_fp)
        fp.parent.mkdir(parents=True, exist_ok=True)
        fp.write_text(body, encoding="utf-8")

        return {
            "ok": True,
            "file_path": str(fp.resolve()),
            "total_sections_processed": len(sections),
        }
    except Exception as e:
        return _err(str(e))


# ---------------------------------------------------------------------------
# 分发与注册表（供 loader / native_tools 挂载）
# ---------------------------------------------------------------------------

_UTIL_HANDLERS: dict[str, Any] = {
    "util:datetime_calc": run_datetime_calc,
    "util:cron_explain": run_cron_explain,
    "util:precise_math": run_precise_math,
    "util:uuid_gen": run_uuid_gen,
    "util:hash_crypto": run_hash_crypto,
    "util:json_jq": run_json_jq,
    "util:regex_test": run_regex_test,
    "util:http_ping": run_http_ping,
    "util:stealth_extract": run_stealth_extract,
    "util:dns_lookup": run_dns_lookup,
    "util:get_weather_lite": run_get_weather_lite,
    "util:ab_test_calc": run_ab_test_calc,
    "util:fake_data_gen": run_fake_data_gen,
    "util:text_diff": run_text_diff,
    "util:funnel_calc": run_funnel_calc,
    "util:generate_office_doc": run_generate_office_doc,
    "util:compose_long_document": run_compose_long_document,
    "util:desktop_message_box": run_desktop_message_box,
    "util:schedule_desktop_reminder": run_schedule_desktop_reminder,
    "util:lark_send_text": run_lark_send_text,
    "util:lark_search_user": run_lark_search_user,
    "util:lark_resolve_user": run_lark_resolve_user,
    "sys:health_stats": run_health_stats,
    "sys:list_env_safe": run_list_env_safe,
}


def dispatch_util_native_tool(tool_id: str, **kwargs: Any) -> dict[str, Any]:
    """
    由 core.native_tools.dispatch_native_tool 转发。
    始终返回 dict（ok/result 或 ok/error），便于 JSON 输出。
    """
    tid = (tool_id or "").strip()
    fn = _UTIL_HANDLERS.get(tid)
    if fn is None:
        return {"ok": False, "error": f"未知 util/sys 工具: {tool_id}"}
    try:
        out = fn(**kwargs)
        if isinstance(out, dict) and "ok" in out:
            return out
        return {"ok": True, "result": out}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def util_tool_ids() -> frozenset[str]:
    return frozenset(_UTIL_HANDLERS.keys())


# 与 loader.NATIVE_TOOLS 条目形状一致：id / label / desc / params
UTIL_TOOLS_NATIVES_LIST: list[dict[str, Any]] = [
    {
        "id": "util:datetime_calc",
        "label": "util:datetime_calc",
        "desc": "日期时间计算。JSON：add_days（int）；可选 base_time（ISO8601）、target_timezone（IANA，默认 Asia/Shanghai）",
        "params": ["add_days"],
    },
    {
        "id": "util:cron_explain",
        "label": "util:cron_explain",
        "desc": "解析 Cron（5 段默认用 APScheduler；6 段含秒需可选安装 croniter）。JSON：cron_expr",
        "params": ["cron_expr"],
    },
    {
        "id": "util:precise_math",
        "label": "util:precise_math",
        "desc": "安全四则运算（Decimal），禁止 eval。JSON：expression（如 \"1024.56 * 3.14 / 2\"）",
        "params": ["expression"],
    },
    {
        "id": "util:uuid_gen",
        "label": "util:uuid_gen",
        "desc": "生成 UUID v4。Action Input 可为 {} 或空",
        "params": [],
    },
    {
        "id": "util:hash_crypto",
        "label": "util:hash_crypto",
        "desc": "哈希/编解码。JSON：text；algo 可选 md5|sha256|base64_encode|base64_decode",
        "params": ["text"],
    },
    {
        "id": "util:json_jq",
        "label": "util:json_jq",
        "desc": "从 JSON 字符串按路径取值（点号与 [0] 下标）。JSON：json_string, path",
        "params": ["json_string", "path"],
    },
    {
        "id": "util:regex_test",
        "label": "util:regex_test",
        "desc": "正则测试。JSON：pattern, test_cases（字符串数组）",
        "params": ["pattern", "test_cases"],
    },
    {
        "id": "util:http_ping",
        "label": "util:http_ping",
        "desc": "HTTP HEAD/GET 探测，超时 3s。JSON：url；可选 method（HEAD|GET）",
        "params": ["url"],
    },
    {
        "id": "util:stealth_extract",
        "label": "util:stealth_extract",
        "desc": "智能路由抓取：默认轻装（curl_cffi / Scrapling Fetcher，via=in_process_fast）；遇 403/503 或常见 WAF 页再调用重装旁路（需 tools/scrapling-service，via=sidecar_heavy）。"
        "旁路未启动时健康检查约 1.5s 内失败即提示启动 uvicorn。JSON：url。",
        "params": ["url"],
    },
    {
        "id": "util:dns_lookup",
        "label": "util:dns_lookup",
        "desc": "DNS 解析。JSON：domain",
        "params": ["domain"],
    },
    {
        "id": "util:get_weather_lite",
        "label": "util:get_weather_lite",
        "desc": "极简天气（免 Key：先 wttr.in，失败则 Open-Meteo；中马等城市有英文别名）。"
        "用户话里若出现城市/地区名须传入 city 或 location；勿传空 JSON。",
        "params": [],
    },
    {
        "id": "util:ab_test_calc",
        "label": "util:ab_test_calc",
        "desc": "A/B 转化率 Z 检验（双尾 P-value，α=0.05）。JSON：visitors_a, conversions_a, visitors_b, conversions_b",
        "params": ["visitors_a", "conversions_a", "visitors_b", "conversions_b"],
    },
    {
        "id": "util:fake_data_gen",
        "label": "util:fake_data_gen",
        "desc": "Faker 生成占位数据（需 pip install faker）。JSON：可选 locale、count（1–50）、fields（name|phone|email|address|company|job）",
        "params": [],
    },
    {
        "id": "util:text_diff",
        "label": "util:text_diff",
        "desc": "两段文本 unified_diff，仅返回 +/- 差异行。JSON：text1（旧）, text2（新）",
        "params": ["text1", "text2"],
    },
    {
        "id": "util:funnel_calc",
        "label": "util:funnel_calc",
        "desc": "漏斗各层人数与可选 ROI。JSON：initial_traffic, conversion_rates；可选 cac、arpu（需同时给）",
        "params": ["initial_traffic", "conversion_rates"],
    },
    {
        "id": "util:generate_office_doc",
        "label": "util:generate_office_doc",
        "desc": "【强制】生成原生 Word/Excel；**绝对禁止**用 core:fs_write 写 .docx/.xlsx。"
        "参数：file_format（docx|xlsx）、file_path、content_json。"
        "docx：content_json.blocks[]，type 为 h1|h2|h3|p|bullet|table（table 用 data 二维数组）。"
        "xlsx：content_json.sheets[]，每项 sheet_name + data 二维数组。"
        "【路径特权】file_path 可为 workspace 相对路径，或 ~/Desktop/、~/Downloads/、~/Documents/ 下的绝对路径（与 core:fs_write 白名单一致）。",
        "params": ["file_format", "file_path", "content_json"],
    },
    {
        "id": "util:compose_long_document",
        "label": "util:compose_long_document",
        "desc": "用于生成极度丰富、字数极多（万字级）的深度报告或长篇文档。不要使用 fs_write，而是将你构思好的文章主题和详细大纲（拆分成 5-10 个具体节点）作为数组传入本工具。"
        "本工具会在后台利用独立线程逐章撰写并自动拼装成完整文件。"
        "JSON：file_path、topic、outline_sections（字符串数组）。路径与 core:fs_write 白名单一致。",
        "params": ["file_path", "topic", "outline_sections"],
    },
    {
        "id": "sys:health_stats",
        "label": "sys:health_stats",
        "desc": "本机 CPU/内存/磁盘余量（psutil）。Action Input 可为 {}",
        "params": [],
    },
    {
        "id": "sys:list_env_safe",
        "label": "sys:list_env_safe",
        "desc": "列出环境变量名（不含值，防泄密）。Action Input 可为 {}",
        "params": [],
    },
    {
        "id": "util:desktop_message_box",
        "label": "util:desktop_message_box",
        "desc": "本机立即弹出可见提醒（Windows 消息框 / macOS 通知 / Linux notify-send）。"
        "JSON：title（可选，默认 Jachin）, message（必填）。"
        "仅**当下**弹出；「某时刻再弹」需系统闹钟/计划任务或到时由调度再调本工具；**禁止**对用户谎称无法弹窗。",
        "params": ["message"],
    },
    {
        "id": "util:schedule_desktop_reminder",
        "label": "util:schedule_desktop_reminder",
        "desc": "【Jachin 桌面】注册定时右下角哨兵提醒（HTTP 127.0.0.1:8002，桌面端须运行）。"
        "JSON：body（必填）, title（可选）；时刻三选一 fire_at_unix_ms | delay_seconds | fire_at_iso（ISO）。"
        "与 util:desktop_message_box（立即弹窗）不同。",
        "params": ["body"],
    },
    {
        "id": "util:lark_send_text",
        "label": "util:lark_send_text",
        "desc": "【飞书/Lark】向指定会话发送纯文本（Open API，需 LARK_APP_ID/SECRET；默认 LARK_CHAT_ID）。"
        "网页摘要场景：先 util:stealth_extract 取正文，再本工具发摘要。**勿**在工具参数中写密钥。"
        "**禁止**把「人名、昵称」填进 chat_id；须 oc_/ou_ 或邮箱/手机，或先 **util:lark_search_user** / util:lark_resolve_user。",
        "params": ["text"],
    },
    {
        "id": "util:lark_search_user",
        "label": "util:lark_search_user",
        "desc": "【飞书】按姓名/昵称搜索用户，返回候选 open_id（多结果须让用户选）。发消息前若仅有姓名须先调本工具或别名表。",
        "params": ["query"],
    },
    {
        "id": "util:lark_resolve_user",
        "label": "util:lark_resolve_user",
        "desc": "【飞书】用邮箱或手机号解析用户 open_id（batch_get_id）；display_name 仅当 LARK_DISPLAY_NAME_MAP 已映射。姓名无映射时优先 util:lark_search_user。",
        "params": ["email"],
    },
]


def _schema_obj(
    props: dict[str, Any],
    required: list[str] | None = None,
) -> dict[str, Any]:
    return {
        "type": "object",
        "properties": props,
        "required": required or [],
    }


# 供上游或 OpenAPI 生成器使用：完整 Schema
UTIL_TOOLS_REGISTRY: dict[str, dict[str, Any]] = {
    "util:datetime_calc": {
        "name": "util:datetime_calc",
        "description": "在指定时区上基于 base_time 增加 add_days，返回 ISO 与时间戳",
        "inputSchema": _schema_obj(
            {
                "base_time": {"type": "string", "description": "ISO8601，缺省为当前时刻"},
                "add_days": {"type": "integer", "description": "要增加的天数，可为负"},
                "target_timezone": {"type": "string", "description": "IANA 时区，默认 Asia/Shanghai"},
            },
            required=[],
        ),
    },
    "util:cron_explain": {
        "name": "util:cron_explain",
        "description": "解释 Cron 并给出下三次触发（UTC ISO）。5 段无需 croniter；6 段含秒需可选安装 croniter",
        "inputSchema": _schema_obj(
            {"cron_expr": {"type": "string", "description": "标准 5 或 6 段 Cron"}},
            required=["cron_expr"],
        ),
    },
    "util:precise_math": {
        "name": "util:precise_math",
        "description": "安全算术表达式，仅 + - * / 与括号、一元正负",
        "inputSchema": _schema_obj(
            {"expression": {"type": "string"}},
            required=["expression"],
        ),
    },
    "util:uuid_gen": {
        "name": "util:uuid_gen",
        "description": "生成随机 UUID4",
        "inputSchema": _schema_obj({}),
    },
    "util:hash_crypto": {
        "name": "util:hash_crypto",
        "description": "MD5/SHA256/Base64",
        "inputSchema": _schema_obj(
            {
                "text": {"type": "string"},
                "algo": {
                    "type": "string",
                    "enum": ["md5", "sha256", "base64_encode", "base64_decode"],
                },
            },
            required=["text"],
        ),
    },
    "util:json_jq": {
        "name": "util:json_jq",
        "description": "轻量 JSON 路径（点号与 [n]）",
        "inputSchema": _schema_obj(
            {"json_string": {"type": "string"}, "path": {"type": "string"}},
            required=["json_string", "path"],
        ),
    },
    "util:regex_test": {
        "name": "util:regex_test",
        "description": "对多段文本做正则匹配与分组提取",
        "inputSchema": _schema_obj(
            {
                "pattern": {"type": "string"},
                "test_cases": {"type": "array", "items": {"type": "string"}},
            },
            required=["pattern", "test_cases"],
        ),
    },
    "util:http_ping": {
        "name": "util:http_ping",
        "description": "HTTP 探测延迟与状态码",
        "inputSchema": _schema_obj(
            {
                "url": {"type": "string"},
                "method": {"type": "string", "enum": ["HEAD", "GET"]},
            },
            required=["url"],
        ),
    },
    "util:stealth_extract": {
        "name": "util:stealth_extract",
        "description": "轻装 curl_cffi 优先；拦截或 403/503 时再调用重装 FastAPI 旁路",
        "inputSchema": _schema_obj(
            {"url": {"type": "string", "description": "目标页面 URL（http/https）"}},
            required=["url"],
        ),
    },
    "util:dns_lookup": {
        "name": "util:dns_lookup",
        "description": "域名解析为 IP 列表",
        "inputSchema": _schema_obj({"domain": {"type": "string"}}, required=["domain"]),
    },
    "util:get_weather_lite": {
        "name": "util:get_weather_lite",
        "description": "极简天气（wttr/Open-Meteo；支持中马等）。用户提到城市时必传 city 或 location",
        "inputSchema": _schema_obj(
            {"city": {"type": "string"}, "location": {"type": "string"}},
        ),
    },
    "util:ab_test_calc": {
        "name": "util:ab_test_calc",
        "description": "两组转化率差异 Z 检验（合并方差），双尾 p-value，报告 is_significant（α=0.05）",
        "inputSchema": _schema_obj(
            {
                "visitors_a": {"type": "integer"},
                "conversions_a": {"type": "integer"},
                "visitors_b": {"type": "integer"},
                "conversions_b": {"type": "integer"},
            },
            required=["visitors_a", "conversions_a", "visitors_b", "conversions_b"],
        ),
    },
    "util:fake_data_gen": {
        "name": "util:fake_data_gen",
        "description": "Faker 占位数据（需安装 faker 包）",
        "inputSchema": _schema_obj(
            {
                "locale": {"type": "string", "description": "BCP 47 / Faker locale，默认 zh_CN"},
                "count": {"type": "integer", "description": "条数 1–50，默认 5"},
                "fields": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "name|phone|email|address|company|job，缺省为全部",
                },
            },
        ),
    },
    "util:text_diff": {
        "name": "util:text_diff",
        "description": "文本/配置行级差异（仅 +/- 行）",
        "inputSchema": _schema_obj(
            {"text1": {"type": "string"}, "text2": {"type": "string"}},
            required=["text1", "text2"],
        ),
    },
    "util:funnel_calc": {
        "name": "util:funnel_calc",
        "description": "漏斗各层人数；可选 CAC×流量与末层 ARPU 算成本、收入、ROI、CPA",
        "inputSchema": _schema_obj(
            {
                "initial_traffic": {"type": "integer"},
                "conversion_rates": {"type": "array", "items": {"type": "number"}},
                "cac": {"type": "number", "description": "单访客获客成本"},
                "arpu": {"type": "number", "description": "漏斗末层单用户平均收入"},
            },
            required=["initial_traffic", "conversion_rates"],
        ),
    },
    "sys:health_stats": {
        "name": "sys:health_stats",
        "description": "本机 CPU、可用内存 MB、磁盘剩余 GB",
        "inputSchema": _schema_obj({}),
    },
    "sys:list_env_safe": {
        "name": "sys:list_env_safe",
        "description": "环境变量名列表（无值）",
        "inputSchema": _schema_obj({}),
    },
    "util:generate_office_doc": {
        "name": "util:generate_office_doc",
        "description": (
            "用于生成原生的 Word (.docx) 报告或 Excel (.xlsx) 数据表。"
            "绝对禁止用 core:fs_write（或等价写入）生成 .docx/.xlsx 富文本后缀文件；必须构造符合要求的 JSON（content_json）交给本工具渲染。"
            "【路径特权】可将文件保存到默认 workspace，或用户真实桌面/下载/文档目录，例如 ~/Desktop/文件名.xlsx、~/Downloads/导出.xlsx、~/Documents/报表.docx。"
            "依赖：pip install python-docx openpyxl。"
        ),
        "inputSchema": _schema_obj(
            {
                "file_format": {
                    "type": "string",
                    "enum": ["docx", "xlsx"],
                    "description": "目标格式；兼容旧字段 file_type",
                },
                "file_path": {
                    "type": "string",
                    "description": (
                        "保存路径：相对路径相对于 workspace；或 ~/Desktop、~/Downloads、~/Documents 等白名单绝对路径。"
                        "扩展名须与 file_format 一致。"
                    ),
                },
                "content_json": {
                    "type": "object",
                    "description": (
                        "docx：必须包含 blocks（顺序块数组）。"
                        "xlsx：必须包含 sheets（多工作表；首表会复用 Workbook 默认表并改名）。"
                        "兼容旧字段名 content_data。"
                    ),
                    "properties": {
                        "blocks": {
                            "type": "array",
                            "description": "仅 docx：按顺序渲染",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "type": {
                                        "type": "string",
                                        "enum": ["h1", "h2", "h3", "p", "bullet", "table"],
                                        "description": "块类型：标题/段落/项目符号/表格",
                                    },
                                    "text": {
                                        "type": "string",
                                        "description": "h1、h2、h3、p、bullet 的正文；table 可省略",
                                    },
                                    "data": {
                                        "type": "array",
                                        "items": {"type": "array"},
                                        "description": "仅 type=table：二维行数组，逐格写入单元格",
                                    },
                                },
                                "required": ["type"],
                            },
                        },
                        "sheets": {
                            "type": "array",
                            "description": "仅 xlsx：多个工作表",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "sheet_name": {
                                        "type": "string",
                                        "description": "工作表名（非法字符会被替换；超长截断至 31 字符）",
                                    },
                                    "data": {
                                        "type": "array",
                                        "items": {"type": "array"},
                                        "description": "二维数组；每行 openpyxl.append",
                                    },
                                },
                                "required": ["sheet_name", "data"],
                            },
                        },
                    },
                },
            },
            required=["file_format", "file_path", "content_json"],
        ),
    },
    "util:compose_long_document": {
        "name": "util:compose_long_document",
        "description": (
            "用于生成极度丰富、字数极多（万字级）的深度报告或长篇文档。不要使用 fs_write，而是将你构思好的文章主题和详细大纲（拆分成 5-10 个具体节点）作为数组传入本工具。"
            "本工具会在后台利用独立线程逐章撰写并自动拼装成完整文件。"
        ),
        "inputSchema": _schema_obj(
            {
                "file_path": {
                    "type": "string",
                    "description": (
                        "保存路径：相对路径相对于 Jachin workspace；或 ~/Desktop、~/Downloads、~/Documents 等白名单绝对路径。"
                        "建议 .md 以便存放 Markdown。"
                    ),
                },
                "topic": {
                    "type": "string",
                    "description": "文档总标题 / 主题，将写入各章提示语中。",
                },
                "outline_sections": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "章节节点标题列表（建议 5–10 条，每条对应一次独立生成后再拼接）。",
                    "minItems": 1,
                },
            },
            required=["file_path", "topic", "outline_sections"],
        ),
    },
    "util:desktop_message_box": {
        "name": "util:desktop_message_box",
        "description": "本机立即弹出消息框或桌面通知；定时提醒需闹钟或计划任务在到时触发",
        "inputSchema": _schema_obj(
            {
                "title": {"type": "string", "description": "标题，默认 Jachin"},
                "message": {"type": "string", "description": "正文，必填"},
            },
            required=["message"],
        ),
    },
    "util:schedule_desktop_reminder": {
        "name": "util:schedule_desktop_reminder",
        "description": (
            "向本机 Jachin 桌面（127.0.0.1:8002）注册定时右下角哨兵提醒；"
            "须桌面端已运行。与立即弹窗的 util:desktop_message_box 不同。"
        ),
        "inputSchema": _schema_obj(
            {
                "title": {"type": "string", "description": "标题，默认 Jachin"},
                "body": {"type": "string", "description": "正文，必填（兼容 message）"},
                "message": {"type": "string", "description": "同 body"},
                "fire_at_unix_ms": {"type": "integer", "description": "到点 Unix 毫秒时间戳"},
                "delay_seconds": {"type": "number", "description": "相对延迟秒数"},
                "fire_at_iso": {
                    "type": "string",
                    "description": "ISO8601；无时区则按 Asia/Shanghai 本地解释",
                },
            },
            required=[],
        ),
    },
    "util:lark_send_text": {
        "name": "util:lark_send_text",
        "description": (
            "向飞书/Lark 指定会话发送纯文本（im:message）。凭证来自 LARK_APP_ID/LARK_APP_SECRET 或 im_channels.yaml；"
            "目标会话默认 LARK_CHAT_ID，可覆盖传入 chat_id。"
            "**禁止**将人名/昵称当作 chat_id；须 oc_/ou_、邮箱、手机，或先 util:lark_search_user / util:lark_resolve_user。"
        ),
        "inputSchema": _schema_obj(
            {
                "text": {"type": "string", "description": "要发送的正文（摘要可直接放这里）"},
                "message": {"type": "string", "description": "同 text"},
                "chat_id": {
                    "type": "string",
                    "description": "receive_id；须 oc_ 会话、ou_ 用户、或邮箱/手机（自动解析）。**勿填人名**（仅有姓名时先 util:lark_search_user）。缺省读 LARK_CHAT_ID / LARK_USER_OPEN_ID",
                },
                "receive_id": {"type": "string", "description": "同 chat_id"},
                "receive_id_type": {
                    "type": "string",
                    "enum": ["chat_id", "open_id", "user_id", "union_id", "email"],
                    "description": "与飞书 API 一致；私聊推荐 open_id（ou_…）",
                },
            },
            required=["text"],
        ),
    },
    "util:lark_search_user": {
        "name": "util:lark_search_user",
        "description": (
            "飞书：按姓名/昵称搜索用户，返回候选 open_id（search/v1/user + 通讯录列表兜底）。"
            "仅姓名时不要直接填 util:lark_send_text；唯一候选可用其 open_id，多条须在回复中列出并请用户确认。"
            "可选配置 LARK_FEISHU_USER_ACCESS_TOKEN 以提升搜索命中率。"
        ),
        "inputSchema": _schema_obj(
            {
                "query": {"type": "string", "description": "搜索关键词（中文名/英文名/昵称）"},
                "name": {"type": "string", "description": "同 query"},
            },
            required=[],
        ),
    },
    "util:lark_resolve_user": {
        "name": "util:lark_resolve_user",
        "description": (
            "飞书通讯录：用邮箱或手机号查询用户 open_id（batch_get_id）。"
            "也可传 display_name（须先在 LARK_DISPLAY_NAME_MAP 配置「英文名→ou_」）。"
            "仅有姓名且未映射时优先 util:lark_search_user。"
            "权限报错时会自动重试一次并提示「版本发布」说明。"
        ),
        "inputSchema": _schema_obj(
            {
                "email": {"type": "string", "description": "飞书企业邮箱"},
                "mobile": {"type": "string", "description": "大陆 11 位手机号"},
                "display_name": {
                    "type": "string",
                    "description": "英文名/昵称；仅当已在 LARK_DISPLAY_NAME_MAP 映射到 ou_/邮箱 时有效",
                },
            },
            required=[],
        ),
    },
}